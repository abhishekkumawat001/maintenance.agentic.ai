import asyncio
import json
import logging
import sqlite3
import random
import os
import cv2
import numpy as np
import base64
import mimetypes
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from enum import Enum
import google.generativeai as genai
from google.generativeai.generative_models import GenerativeModel
from dotenv import load_dotenv
import PyPDF2
import docx
from pathlib import Path
import hashlib

load_dotenv()

class LLMConfig:
    """Configuration for LLM providers"""
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "your-gemini-api-key")

gemini_key = os.getenv("GEMINI_API_KEY")
print(f"Gemini API Key status: {'Configured' if gemini_key else 'Not configured'}")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enhanced Data Models
@dataclass
class SensorData:
    sensor_id: str
    timestamp: datetime
    temperature: float
    vibration: float
    pressure: float
    humidity: float
    sound_level: float

@dataclass
class MachineDocument:
    doc_id: str
    machine_id: str
    file_path: str
    file_type: str
    content: str
    upload_date: datetime
    summary: str = ""

@dataclass
class VisualInspection:
    inspection_id: str
    machine_id: str
    file_path: str
    file_type: str  # image or video
    defects_detected: List[Dict]
    timestamp: datetime
    ai_analysis: str = ""

@dataclass
class FaultDiagnosis:
    fault_id: str
    machine_id: str
    severity: str
    confidence: float
    description: str
    root_cause: str
    recommended_actions: List[str]
    estimated_downtime: int
    supporting_evidence: List[str]  # References to documents/images
    conversation_history: List[Dict] = None

@dataclass
class DiagnosisSession:
    session_id: str
    machine_id: str
    start_time: datetime
    current_diagnosis: Optional[FaultDiagnosis]
    conversation_history: List[Dict]
    uploaded_files: List[str]
    status: str  # active, resolved, escalated

class MaintenanceType(Enum):
    PREVENTIVE = "preventive"
    CORRECTIVE = "corrective"
    PREDICTIVE = "predictive"
    EMERGENCY = "emergency"

@dataclass
class MaintenanceTask:
    task_id: str
    equipment_id: str
    task_type: MaintenanceType
    priority: int
    description: str
    scheduled_date: datetime
    estimated_duration: int
    required_parts: List[str]
    assigned_technician: Optional[str] = None

# Enhanced Document Processing
class DocumentProcessor:
    def __init__(self, storage_dir: str = "uploads"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
    def save_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded file and return new path"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        
        # Copy file to storage directory
        import shutil
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def process_document(self, file_path: str, machine_id: str) -> MachineDocument:
        """Process uploaded document and extract content"""
        file_path = Path(file_path)
        file_type = file_path.suffix.lower()
        
        # Extract text based on file type
        if file_type == '.pdf':
            content = self.extract_text_from_pdf(str(file_path))
        elif file_type == '.docx':
            content = self.extract_text_from_docx(str(file_path))
        elif file_type in ['.txt', '.md']:
            content = self.extract_text_from_txt(str(file_path))
        else:
            content = f"Unsupported file type: {file_type}"
        
        # Generate document ID
        doc_id = hashlib.md5(f"{machine_id}_{file_path.name}_{datetime.now()}".encode()).hexdigest()[:12]
        
        return MachineDocument(
            doc_id=doc_id,
            machine_id=machine_id,
            file_path=str(file_path),
            file_type=file_type,
            content=content,
            upload_date=datetime.now()
        )

# Enhanced Visual Processing
class EnhancedVisionProcessor:
    def __init__(self, llm_provider, storage_dir: str = "visual_uploads"):
        self.llm_provider = llm_provider
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
    
    def save_visual_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded visual file"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        
        import shutil
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def analyze_image(self, image_path: str) -> Dict[str, Any]:
        """Enhanced image analysis with OpenCV"""
        try:
            image = cv2.imread(image_path)
            if image is None:
                return {"error": "Could not load image"}
            
            # Multiple analysis techniques
            analysis_results = {}
            
            # 1. Basic defect detection
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            edges = cv2.Canny(gray, 50, 150)
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            defects = []
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > 500:  # Threshold for significant defects
                    x, y, w, h = cv2.boundingRect(contour)
                    defects.append({
                        'location': {'x': int(x), 'y': int(y), 'width': int(w), 'height': int(h)},
                        'area': float(area),
                        'type': 'potential_defect'
                    })
            
            # 2. Corrosion detection (rust-like colors)
            hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
            rust_lower = np.array([0, 50, 50])
            rust_upper = np.array([20, 255, 255])
            rust_mask = cv2.inRange(hsv, rust_lower, rust_upper)
            rust_area = cv2.countNonZero(rust_mask)
            
            # 3. Crack detection using morphological operations
            kernel = np.ones((3,3), np.uint8)
            morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
            crack_edges = cv2.Canny(morph, 100, 200)
            crack_lines = cv2.HoughLinesP(crack_edges, 1, np.pi/180, threshold=50, minLineLength=30, maxLineGap=10)
            
            analysis_results = {
                'defects_found': len(defects),
                'defects': defects,
                'rust_area_pixels': int(rust_area),
                'potential_cracks': len(crack_lines) if crack_lines is not None else 0,
                'image_dimensions': {'width': image.shape[1], 'height': image.shape[0]},
                'analysis_timestamp': datetime.now().isoformat()
            }
            
            return analysis_results
            
        except Exception as e:
            logger.error(f"Vision processing error: {e}")
            return {"error": str(e)}
    
    def analyze_video(self, video_path: str) -> Dict[str, Any]:
        """Basic video analysis for defect detection"""
        try:
            cap = cv2.VideoCapture(video_path)
            if not cap.isOpened():
                return {"error": "Could not open video"}
            
            frame_count = 0
            total_defects = 0
            frame_analyses = []
            
            # Analyze every 30th frame to avoid processing every frame
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                
                if frame_count % 30 == 0:  # Process every 30th frame
                    # Save frame temporarily
                    temp_frame_path = f"temp_frame_{frame_count}.jpg"
                    cv2.imwrite(temp_frame_path, frame)
                    
                    # Analyze frame
                    frame_analysis = self.analyze_image(temp_frame_path)
                    if 'defects_found' in frame_analysis:
                        total_defects += frame_analysis['defects_found']
                        frame_analyses.append({
                            'frame_number': frame_count,
                            'defects': frame_analysis['defects_found'],
                            'timestamp': frame_count / cap.get(cv2.CAP_PROP_FPS) if cap.get(cv2.CAP_PROP_FPS) > 0 else 0
                        })
                    
                    # Clean up temp file
                    if os.path.exists(temp_frame_path):
                        os.remove(temp_frame_path)
                
                frame_count += 1
            
            cap.release()
            
            return {
                'total_frames': frame_count,
                'frames_analyzed': len(frame_analyses),
                'total_defects_detected': total_defects,
                'frame_analyses': frame_analyses,
                'analysis_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Video processing error: {e}")
            return {"error": str(e)}
    
    async def analyze_with_gemini_vision(self, file_path: str, machine_context: str = "") -> str:
        """Analyze visual content with Gemini Pro Vision"""
        try:
            # Prepare image for Gemini
            with open(file_path, 'rb') as f:
                image_data = f.read()
            
            # Create image part for Gemini
            image_part = {
                "mime_type": mimetypes.guess_type(file_path)[0] or "image/jpeg",
                "data": base64.b64encode(image_data).decode()
            }
            
            prompt = f"""
            As an expert industrial maintenance engineer, analyze this image of industrial equipment:
            
            Machine Context: {machine_context}
            
            Please provide a detailed analysis focusing on:
            1. Visible defects, damage, or wear
            2. Corrosion, rust, or material degradation
            3. Misalignments or structural issues
            4. Leaks, cracks, or surface irregularities
            5. Overall condition assessment
            6. Immediate safety concerns
            7. Recommended actions based on visual inspection
            
            Provide specific locations and severity levels for any issues found.
            Format your response with clear headings and actionable recommendations.
            """
            
            # For now, use text-only Gemini since vision may not be available
            # In production, you would use Gemini Pro Vision
            response = await self.llm_provider.get_best_response(
                f"{prompt}\n\nNote: Visual analysis completed using computer vision algorithms. "
                f"Please provide maintenance recommendations based on typical defects found in industrial equipment imagery."
            )
            
            return response
            
        except Exception as e:
            logger.error(f"Gemini vision analysis error: {e}")
            return f"Visual analysis error: {e}"

# Enhanced LLM Provider
class EnhancedLLMProvider:
    def __init__(self):
        self.init_gemini()
    
    def init_gemini(self):
        """Initialize Gemini provider"""
        try:
            if LLMConfig.GEMINI_API_KEY and LLMConfig.GEMINI_API_KEY != "your-gemini-api-key":
                genai.configure(api_key=LLMConfig.GEMINI_API_KEY)
                self.gemini_model = genai.GenerativeModel('gemini-1.5-flash')
                logger.info("Gemini model initialized successfully")
            else:
                self.gemini_model = None
                logger.warning("Gemini API key not configured")
        except Exception as e:
            logger.warning(f"Error initializing Gemini: {e}")
            self.gemini_model = None
    
    async def call_gemini(self, prompt: str) -> str:
        """Call Gemini API"""
        try:
            if not self.gemini_model:
                return "Gemini API not configured. Please set GEMINI_API_KEY environment variable."
                
            response = self.gemini_model.generate_content(prompt)
            return response.text
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            return f"Error calling Gemini: {e}"
    
    async def get_best_response(self, prompt: str) -> str:
        """Get response from Gemini"""
        return await self.call_gemini(prompt)
    
    async def analyze_with_context(self, prompt: str, context_docs: List[MachineDocument], visual_data: List[VisualInspection] = None) -> str:
        """Analyze with machine-specific context"""
        enhanced_prompt = f"""
        MACHINE-SPECIFIC CONTEXT:
        {self._format_document_context(context_docs)}
        
        VISUAL INSPECTION DATA:
        {self._format_visual_context(visual_data) if visual_data else "No visual data available"}
        
        MAINTENANCE EXPERT ANALYSIS REQUEST:
        {prompt}
        
        Please provide analysis based on the specific machine documentation and visual inspection data provided above.
        Reference specific sections of the documentation when relevant.
        """
        
        return await self.get_best_response(enhanced_prompt)
    
    def _format_document_context(self, docs: List[MachineDocument]) -> str:
        """Format document context for LLM"""
        if not docs:
            return "No machine documentation available"
        
        context = ""
        for doc in docs:
            context += f"""
            Document: {doc.file_path}
            Type: {doc.file_type}
            Content Summary: {doc.content[:1000]}...
            ---
            """
        return context
    
    def _format_visual_context(self, visual_data: List[VisualInspection]) -> str:
        """Format visual inspection context"""
        if not visual_data:
            return "No visual inspection data available"
        
        context = ""
        for inspection in visual_data:
            context += f"""
            Visual Inspection: {inspection.file_path}
            Type: {inspection.file_type}
            Defects Detected: {len(inspection.defects_detected)}
            AI Analysis: {inspection.ai_analysis[:500]}...
            ---
            """
        return context
    
    async def close(self):
        """Cleanup resources"""
        pass

# Enhanced Knowledge Base
class EnhancedMaintenanceKnowledgeBase:
    def __init__(self, db_path: str = "enhanced_maintenance.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Initialize enhanced SQLite database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Equipment table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS equipment (
                id TEXT PRIMARY KEY,
                name TEXT,
                type TEXT,
                model TEXT,
                manufacturer TEXT,
                installation_date DATE,
                last_maintenance DATE
            )
        ''')
        
        # Machine documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS machine_documents (
                doc_id TEXT PRIMARY KEY,
                machine_id TEXT,
                file_path TEXT,
                file_type TEXT,
                content TEXT,
                summary TEXT,
                upload_date TIMESTAMP,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Visual inspections table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS visual_inspections (
                inspection_id TEXT PRIMARY KEY,
                machine_id TEXT,
                file_path TEXT,
                file_type TEXT,
                defects_data TEXT,
                ai_analysis TEXT,
                timestamp TIMESTAMP,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Diagnosis sessions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS diagnosis_sessions (
                session_id TEXT PRIMARY KEY,
                machine_id TEXT,
                start_time TIMESTAMP,
                current_diagnosis TEXT,
                conversation_history TEXT,
                uploaded_files TEXT,
                status TEXT,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Enhanced maintenance history
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS maintenance_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                equipment_id TEXT,
                maintenance_date DATE,
                type TEXT,
                description TEXT,
                technician TEXT,
                cost REAL,
                downtime_hours INTEGER,
                supporting_documents TEXT,
                FOREIGN KEY (equipment_id) REFERENCES equipment (id)
            )
        ''')
        
        # Fault patterns table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS fault_patterns (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                equipment_type TEXT,
                symptoms TEXT,
                root_cause TEXT,
                solution TEXT,
                confidence_score REAL,
                supporting_evidence TEXT
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def add_machine_document(self, document: MachineDocument) -> bool:
        """Add machine document to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO machine_documents 
                (doc_id, machine_id, file_path, file_type, content, summary, upload_date)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                document.doc_id,
                document.machine_id,
                document.file_path,
                document.file_type,
                document.content,
                document.summary,
                document.upload_date
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding machine document: {e}")
            return False
    
    def get_machine_documents(self, machine_id: str) -> List[MachineDocument]:
        """Get all documents for a specific machine"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM machine_documents WHERE machine_id = ?
        ''', (machine_id,))
        
        results = cursor.fetchall()
        conn.close()
        
        documents = []
        for row in results:
            doc = MachineDocument(
                doc_id=row[0],
                machine_id=row[1],
                file_path=row[2],
                file_type=row[3],
                content=row[4],
                upload_date=datetime.fromisoformat(row[6]),
                summary=row[5] or ""
            )
            documents.append(doc)
        
        return documents
    
    def add_visual_inspection(self, inspection: VisualInspection) -> bool:
        """Add visual inspection to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO visual_inspections 
                (inspection_id, machine_id, file_path, file_type, defects_data, ai_analysis, timestamp)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                inspection.inspection_id,
                inspection.machine_id,
                inspection.file_path,
                inspection.file_type,
                json.dumps(inspection.defects_detected),
                inspection.ai_analysis,
                inspection.timestamp
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding visual inspection: {e}")
            return False
    
    def get_visual_inspections(self, machine_id: str) -> List[VisualInspection]:
        """Get all visual inspections for a machine"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM visual_inspections WHERE machine_id = ?
        ''', (machine_id,))
        
        results = cursor.fetchall()
        conn.close()
        
        inspections = []
        for row in results:
            inspection = VisualInspection(
                inspection_id=row[0],
                machine_id=row[1],
                file_path=row[2],
                file_type=row[3],
                defects_detected=json.loads(row[4]) if row[4] else [],
                timestamp=datetime.fromisoformat(row[6]),
                ai_analysis=row[5] or ""
            )
            inspections.append(inspection)
        
        return inspections
    
    def save_diagnosis_session(self, session: DiagnosisSession) -> bool:
        """Save diagnosis session"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO diagnosis_sessions 
                (session_id, machine_id, start_time, current_diagnosis, conversation_history, uploaded_files, status)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                session.session_id,
                session.machine_id,
                session.start_time,
                json.dumps(session.current_diagnosis.__dict__ if session.current_diagnosis else {}),
                json.dumps(session.conversation_history),
                json.dumps(session.uploaded_files),
                session.status
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error saving diagnosis session: {e}")
            return False
    
    # ... (include other methods from original knowledge base)
    def add_equipment(self, equipment_data: Dict) -> bool:
        """Add new equipment to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO equipment (id, name, type, model, manufacturer, installation_date, last_maintenance)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                equipment_data['id'],
                equipment_data['name'],
                equipment_data['type'],
                equipment_data.get('model', 'Unknown'),
                equipment_data.get('manufacturer', 'Unknown'),
                equipment_data.get('installation_date', datetime.now().strftime('%Y-%m-%d')),
                equipment_data.get('last_maintenance', 'Never')
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding equipment: {e}")
            return False
    
    def get_all_equipment(self) -> List[Dict]:
        """Get all equipment from database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM equipment')
        equipment = cursor.fetchall()
        columns = [description[0] for description in cursor.description]
        conn.close()
        
        return [dict(zip(columns, row)) for row in equipment]

# Enhanced Diagnostic Engine
class EnhancedDiagnosticEngine:
    def __init__(self, knowledge_base: EnhancedMaintenanceKnowledgeBase, llm_provider: EnhancedLLMProvider):
        self.knowledge_base = knowledge_base
        self.llm_provider = llm_provider
    
    async def comprehensive_diagnosis(self, machine_id: str, sensor_data: Dict = None, visual_data: List[VisualInspection] = None, user_description: str = "") -> FaultDiagnosis:
        """Perform comprehensive diagnosis using all available data"""
        
        # Get machine documents
        machine_docs = self.knowledge_base.get_machine_documents(machine_id)
        
        # Get visual inspections if not provided
        if visual_data is None:
            visual_data = self.knowledge_base.get_visual_inspections(machine_id)
        
        # Prepare comprehensive analysis prompt
        prompt = f"""
        COMPREHENSIVE MACHINE DIAGNOSIS REQUEST
        
        Machine ID: {machine_id}
        User Description: {user_description}
        
        AVAILABLE DATA:
        1. Machine Documentation: {len(machine_docs)} documents available
        2. Visual Inspections: {len(visual_data)} inspections available
        3. Sensor Data: {"Available" if sensor_data else "Not available"}
        
        ANALYSIS REQUEST:
        Based on the machine-specific documentation and visual inspection data, please provide:
        
        1. PRIMARY DIAGNOSIS:
           - Most likely root cause
           - Confidence level (0.0-1.0)
           - Severity assessment (low/medium/high/critical)
        
        2. SUPPORTING EVIDENCE:
           - Reference specific documentation sections
           - Cite visual inspection findings
           - Correlate with sensor anomalies
        
        3. RECOMMENDED DIAGNOSTIC STEPS:
           - Immediate actions to confirm diagnosis
           - Additional tests or inspections needed
           - Safety precautions
        
        4. ALTERNATIVE DIAGNOSES:
           - Other possible causes
           - How to differentiate between them
        
        5. REPAIR STRATEGY:
           - Step-by-step repair process
           - Required parts and tools
           - Estimated time and complexity
        
        Please provide specific, actionable recommendations based on the machine's actual documentation and inspection data.
        """
        
        # Get AI analysis with context
        ai_response = await self.llm_provider.analyze_with_context(prompt, machine_docs, visual_data)
        
        # Parse response and create diagnosis
        diagnosis = FaultDiagnosis(
            fault_id=f"DIAG_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            machine_id=machine_id,
            severity="medium",  # Default, should be parsed from AI response
            confidence=0.7,     # Default, should be parsed from AI response
            description=f"AI-assisted comprehensive diagnosis for {machine_id}",
            root_cause="Analysis in progress",
            recommended_actions=["Review AI analysis", "Implement suggested diagnostic steps"],
            estimated_downtime=4,
            supporting_evidence=[doc.file_path for doc in machine_docs] + [vis.file_path for vis in visual_data],
            conversation_history=[{
                "timestamp": datetime.now().isoformat(),
                "type": "initial_diagnosis",
                "content": ai_response
            }]
        )
        
        # Try to parse structured information from AI response
        try:
            # In a real implementation, you might use more sophisticated parsing
            # or prompt the AI to return structured JSON
            if "high" in ai_response.lower() or "critical" in ai_response.lower():
                diagnosis.severity = "high"
            elif "low" in ai_response.lower():
                diagnosis.severity = "low"
            
            # Update root cause and actions based on AI response
            diagnosis.root_cause = ai_response[:200] + "..." if len(ai_response) > 200 else ai_response
            
        except Exception as e:
            logger.error(f"Error parsing AI diagnosis: {e}")
        
        return diagnosis
    
    async def interactive_diagnosis(self, session: DiagnosisSession, user_input: str) -> Tuple[str, DiagnosisSession]:
        """Handle interactive diagnosis conversation"""
        
        # Add user input to conversation history
        session.conversation_history.append({
            "timestamp": datetime.now().isoformat(),
            "speaker": "user",
            "content": user_input
        })
        
        # Get machine context
        machine_docs = self.knowledge_base.get_machine_documents(session.machine_id)
        visual_data = self.knowledge_base.get_visual_inspections(session.machine_id)
        
        # Prepare conversational prompt
        prompt = f"""
        INTERACTIVE DIAGNOSIS SESSION
        
        Machine ID: {session.machine_id}
        Session Status: {session.status}
        
        CONVERSATION HISTORY:
        {self._format_conversation_history(session.conversation_history)}
        
        LATEST USER INPUT: {user_input}
        
        CONTEXT:
        - Current Diagnosis: {session.current_diagnosis.root_cause if session.current_diagnosis else "No diagnosis yet"}
        - Uploaded Files: {len(session.uploaded_files)} files
        
        As an expert maintenance engineer, please:
        1. Respond to the user's latest input