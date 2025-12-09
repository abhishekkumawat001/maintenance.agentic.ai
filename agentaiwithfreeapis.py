import asyncio
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass
from enum import Enum
import aiohttp
from dotenv import load_dotenv
import PyPDF2
import docx
from pathlib import Path
import hashlib
import shutil
load_dotenv()

class LLMConfig:
    """Configuration for free LLM providers"""
    # Free providers only
    GROQ_API_KEY = os.getenv("GROQ_API_KEY", "your-groq-api-key")
    HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY", "your-hf-api-key")
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "your-gemini-api-key")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enhanced Data Models
@dataclass
class SensorData:
    sensor_id: str
    timestamp: datetime
    temperature: float
    vibration: float
    pressure: float
    humidity: float
    sound_level: float

@dataclass
class MachineDocument:
    doc_id: str
    machine_id: str
    file_path: str
    file_type: str
    content: str
    upload_date: datetime
    summary: str = ""

@dataclass
class VisualInspection:
    inspection_id: str
    machine_id: str
    file_path: str
    file_type: str  # image or video
    defects_detected: List[Dict]
    timestamp: datetime
    ai_analysis: str = ""

@dataclass
class FaultDiagnosis:
    fault_id: str
    machine_id: str
    severity: str
    confidence: float
    description: str
    root_cause: str
    recommended_actions: List[str]
    estimated_downtime: int
    supporting_evidence: List[str]  # References to documents/images
    conversation_history: Optional[List[Dict]] = None

@dataclass
class DiagnosisSession:
    session_id: str
    machine_id: str
    start_time: datetime
    current_diagnosis: Optional[FaultDiagnosis]
    conversation_history: List[Dict]
    uploaded_files: List[str]
    status: str  # active, resolved, escalated

class MaintenanceType(Enum):
    PREVENTIVE = "preventive"
    CORRECTIVE = "corrective"
    PREDICTIVE = "predictive"
    EMERGENCY = "emergency"

@dataclass
class MaintenanceTask:
    task_id: str
    equipment_id: str
    task_type: MaintenanceType
    priority: int
    description: str
    scheduled_date: datetime
    estimated_duration: int
    required_parts: List[str]
    assigned_technician: Optional[str] = None

# Enhanced Document Processing
class DocumentProcessor:
    def __init__(self, storage_dir: str = "uploads"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
    def save_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded file and return new path"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def process_document(self, file_path: str, machine_id: str) -> MachineDocument:
        """Process uploaded document and extract content"""
        file_path_obj = Path(file_path)
        file_type = file_path_obj.suffix.lower()
        
        # Extract text based on file type
        if file_type == '.pdf':
            content = self.extract_text_from_pdf(file_path)
        elif file_type == '.docx':
            content = self.extract_text_from_docx(file_path)
        elif file_type in ['.txt', '.md']:
            content = self.extract_text_from_txt(file_path)
        else:
            content = f"Unsupported file type: {file_type}"
        
        # Generate document ID
        doc_id = hashlib.md5(f"{machine_id}_{file_path_obj.name}_{datetime.now()}".encode()).hexdigest()[:12]
        
        return MachineDocument(
            doc_id=doc_id,
            machine_id=machine_id,
            file_path=file_path,
            file_type=file_type,
            content=content,
            upload_date=datetime.now()
        )

class EnhancedLLMProvider:
    def __init__(self):
        self.current_provider = "auto"  # auto, groq, gemini, huggingface, local
        self.init_providers()
    
    def set_provider(self, provider: str):
        """Set specific LLM provider"""
        valid_providers = ["auto", "groq", "gemini", "huggingface", "local"]
        if provider.lower() in valid_providers:
            self.current_provider = provider.lower()
            print(f"LLM provider set to: {self.current_provider}")
        else:
            print(f"Invalid provider. Choose from: {valid_providers}")
    
    def get_current_provider(self):
        """Get current LLM provider"""
        return self.current_provider
    
    def init_providers(self):
        """Initialize free LLM providers only"""
        # Groq (Free tier available)
        self.groq_configured = (LLMConfig.GROQ_API_KEY and 
                              LLMConfig.GROQ_API_KEY != "your-groq-api-key")
        
        # Hugging Face (Free tier available)
        self.hf_configured = (LLMConfig.HUGGINGFACE_API_KEY and 
                            LLMConfig.HUGGINGFACE_API_KEY != "your-hf-api-key")
        
        # Gemini (Free tier available)
        self.gemini_configured = (LLMConfig.GEMINI_API_KEY and 
                                LLMConfig.GEMINI_API_KEY != "your-gemini-api-key")
        
        logger.info(f"Free providers configured: "
                   f"Groq={self.groq_configured}, HF={self.hf_configured}, "
                   f"Gemini={self.gemini_configured}")
    
    async def call_groq(self, prompt: str) -> str:
        """Call Groq API (free tier available)"""
        try:
            if not self.groq_configured:
                return "Groq API not configured"
            
            headers = {
                "Authorization": f"Bearer {LLMConfig.GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": "llama3-8b-8192",  # Free model
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 2048,
                "temperature": 0.7
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        return data["choices"][0]["message"]["content"].strip()
                    else:
                        error_text = await response.text()
                        logger.error(f"Groq API error {response.status}: {error_text}")
                        return f"Groq API error: {response.status}"
                        
        except Exception as e:
            logger.error(f"Groq API error: {e}")
            return f"Error calling Groq: {e}"
    
    async def call_huggingface(self, prompt: str) -> str:
        """Call Hugging Face API (free tier available)"""
        try:
            if not self.hf_configured:
                return "Hugging Face API not configured"
            
            headers = {
                "Authorization": f"Bearer {LLMConfig.HUGGINGFACE_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": 1000,
                    "temperature": 0.7,
                    "return_full_text": False
                }
            }
            
            # Using a free model
            model_url = "https://api-inference.huggingface.co/models/microsoft/DialoGPT-large"
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    model_url,
                    headers=headers,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        if isinstance(data, list) and len(data) > 0:
                            return data[0].get("generated_text", "No response generated")
                        return str(data)
                    else:
                        error_text = await response.text()
                        logger.error(f"HF API error {response.status}: {error_text}")
                        return f"HF API error: {response.status}"
                        
        except Exception as e:
            logger.error(f"HF API error: {e}")
            return f"Error calling Hugging Face: {e}"
    
    async def call_gemini(self, prompt: str) -> str:
        """Call Gemini API (free tier available)"""
        try:
            if not self.gemini_configured:
                return "Gemini API not configured"
            
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={LLMConfig.GEMINI_API_KEY}"
            
            payload = {
                "contents": [{
                    "parts": [{"text": prompt}]
                }],
                "generationConfig": {
                    "temperature": 0.7,
                    "maxOutputTokens": 2048
                }
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    url,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        if 'candidates' in data and len(data['candidates']) > 0:
                            content = data['candidates'][0]['content']['parts'][0]['text']
                            return content.strip()
                        return "No response generated"
                    else:
                        error_text = await response.text()
                        logger.error(f"Gemini API error {response.status}: {error_text}")
                        return f"Gemini API error: {response.status}"
                        
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            return f"Error calling Gemini: {e}"
    
    async def call_local_fallback(self, prompt: str) -> str:
        """Local fallback with pre-defined responses"""
        maintenance_keywords = {
            'bearing': "Common bearing failure causes: inadequate lubrication, misalignment, contamination, overloading, and fatigue. Recommended actions: check lubrication, inspect alignment, replace if worn.",
            'pump': "Common pump issues: cavitation, impeller wear, seal failure, misalignment. Check suction pressure, inspect impeller, verify shaft alignment.",
            'motor': "Common motor problems: overheating, bearing failure, winding damage. Check temperature, vibration levels, electrical connections.",
            'vibration': "Excessive vibration causes: imbalance, misalignment, bearing wear, loose components. Perform vibration analysis, check alignment, inspect bearings.",
            'noise': "Unusual noise sources: bearing wear, cavitation, loose parts, misalignment. Identify noise type and location for proper diagnosis.",
            'leak': "Common leak sources: seal failure, gasket degradation, pipe joints, valve packing. Inspect seals, replace gaskets, check torque specifications."
        }
        
        prompt_lower = prompt.lower()
        for keyword, response in maintenance_keywords.items():
            if keyword in prompt_lower:
                return f"MAINTENANCE ANALYSIS (Local Fallback): {response}"
        
        return "MAINTENANCE ANALYSIS (Local Fallback): Unable to analyze without specific equipment information. Please provide more details about the equipment type, symptoms, and operating conditions."
    
    async def get_best_response(self, prompt: str) -> str:
        """Get response from specified provider or try providers in order"""
        if self.current_provider == "groq" and self.groq_configured:
            return await self.call_groq(prompt)
        elif self.current_provider == "gemini" and self.gemini_configured:
            return await self.call_gemini(prompt)
        elif self.current_provider == "huggingface" and self.hf_configured:
            return await self.call_huggingface(prompt)
        elif self.current_provider == "local":
            return await self.call_local_fallback(prompt)
        elif self.current_provider == "auto":
            # Try providers in order of preference
            providers = [
                ("Groq", self.call_groq) if self.groq_configured else None,
                ("Gemini", self.call_gemini) if self.gemini_configured else None,
                ("HuggingFace", self.call_huggingface) if self.hf_configured else None,
                ("Local Fallback", self.call_local_fallback)
            ]
            
            # Filter out None providers
            providers = [p for p in providers if p is not None]
            
            for provider_name, provider_func in providers:
                try:
                    logger.info(f"Trying {provider_name} provider...")
                    response = await provider_func(prompt)
                    
                    # Check if response indicates an error
                    if not any(error_term in response.lower() for error_term in 
                              ['error', 'quota', 'exceeded', 'not configured']):
                        logger.info(f"Success with {provider_name}")
                        return response
                    else:
                        logger.warning(f"{provider_name} failed: {response[:100]}...")
                        
                except Exception as e:
                    logger.error(f"{provider_name} provider failed: {e}")
                    continue
            
            # If all providers fail, return local fallback
            return await self.call_local_fallback(prompt)
        else:
            return "Invalid provider configuration"

    async def analyze_with_context(self, prompt: str, context_docs: List[MachineDocument], visual_data: Optional[List[VisualInspection]] = None) -> str:
        """Analyze with machine-specific context"""
        enhanced_prompt = f"""
        MACHINE-SPECIFIC CONTEXT:
        {self._format_document_context(context_docs)}
        
        VISUAL INSPECTION DATA:
        {self._format_visual_context(visual_data) if visual_data else "No visual data available"}
        
        MAINTENANCE EXPERT ANALYSIS REQUEST:
        {prompt}
        
        Please provide analysis based on the specific machine documentation and visual inspection data provided above.
        Reference specific sections of the documentation when relevant.
        """
        
        return await self.get_best_response(enhanced_prompt)
    
    def _format_document_context(self, docs: List[MachineDocument]) -> str:
        """Format document context for LLM"""
        if not docs:
            return "No machine documentation available"
        
        context = ""
        for doc in docs:
            context += f"""
            Document: {doc.file_path}
            Type: {doc.file_type}
            Content Summary: {doc.content[:500]}...
            ---
            """
        return context
    
    def _format_visual_context(self, visual_data: List[VisualInspection]) -> str:
        """Format visual inspection context"""
        if not visual_data:
            return "No visual inspection data available"
        
        context = ""
        for inspection in visual_data:
            context += f"""
            Visual Inspection: {inspection.file_path}
            Type: {inspection.file_type}
            Defects Detected: {len(inspection.defects_detected)}
            AI Analysis: {inspection.ai_analysis[:300]}...
            ---
            """
        return context
    
    async def close(self):
        """Cleanup resources"""
        pass

# Enhanced Smart Maintenance System
    async def close(self):
        """Cleanup resources"""
        pass

# Missing classes that need to be implemented
class EnhancedMaintenanceKnowledgeBase:
    def __init__(self):
        self.equipment_data = []
        self.documents = []
        self.visual_inspections = []
    
    def add_equipment(self, equipment_data: Dict) -> bool:
        """Add equipment to knowledge base"""
        try:
            self.equipment_data.append(equipment_data)
            return True
        except Exception:
            return False
    
    def get_all_equipment(self) -> List[Dict]:
        """Get all equipment"""
        return self.equipment_data
    
    def add_machine_document(self, document: MachineDocument) -> bool:
        """Add machine document"""
        try:
            self.documents.append(document)
            return True
        except Exception:
            return False
    
    def get_machine_documents(self, machine_id: str) -> List[MachineDocument]:
        """Get documents for specific machine"""
        return [doc for doc in self.documents if doc.machine_id == machine_id]
    
    def add_visual_inspection(self, inspection: VisualInspection) -> bool:
        """Add visual inspection"""
        try:
            self.visual_inspections.append(inspection)
            return True
        except Exception:
            return False
    
    def get_visual_inspections(self, machine_id: str) -> List[VisualInspection]:
        """Get visual inspections for specific machine"""
        return [vi for vi in self.visual_inspections if vi.machine_id == machine_id]

class EnhancedDiagnosticEngine:
    def __init__(self, knowledge_base: EnhancedMaintenanceKnowledgeBase, llm_provider: EnhancedLLMProvider):
        self.knowledge_base = knowledge_base
        self.llm_provider = llm_provider
        self.active_sessions = {}
    
    async def start_interactive_diagnosis(self, machine_id: str) -> DiagnosisSession:
        """Start new diagnosis session"""
        session_id = f"DIAG_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        session = DiagnosisSession(
            session_id=session_id,
            machine_id=machine_id,
            start_time=datetime.now(),
            current_diagnosis=None,
            conversation_history=[],
            uploaded_files=[],
            status="active"
        )
        self.active_sessions[session_id] = session
        return session
    
    async def comprehensive_diagnosis(self, machine_id: str, user_description: str) -> FaultDiagnosis:
        """Perform comprehensive diagnosis"""
        fault_id = f"FAULT_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Get machine data
        machine_docs = self.knowledge_base.get_machine_documents(machine_id)
        visual_data = self.knowledge_base.get_visual_inspections(machine_id)
        
        # Create diagnosis prompt
        prompt = f"""
        COMPREHENSIVE FAULT DIAGNOSIS
        
        Machine ID: {machine_id}
        Problem Description: {user_description}
        
        Available Documents: {len(machine_docs)}
        Visual Inspections: {len(visual_data)}
        
        Please provide:
        1. Root cause analysis
        2. Severity assessment (low/medium/high/critical)
        3. Confidence level (0-1)
        4. Recommended actions
        5. Estimated downtime in hours
        """
        
        response = await self.llm_provider.analyze_with_context(prompt, machine_docs, visual_data)
        
        # Create diagnosis object
        diagnosis = FaultDiagnosis(
            fault_id=fault_id,
            machine_id=machine_id,
            severity="medium",
            confidence=0.8,
            description=user_description[:200],
            root_cause=response[:300],
            recommended_actions=[
                "Perform detailed inspection",
                "Check maintenance records",
                "Monitor for 24 hours"
            ],
            estimated_downtime=4,
            supporting_evidence=[f"AI Analysis: {response[:100]}"]
        )
        
        return diagnosis
    
    async def continue_diagnosis(self, session_id: str, user_input: str) -> tuple:
        """Continue diagnosis conversation"""
        if session_id not in self.active_sessions:
            raise ValueError("Session not found")
        
        session = self.active_sessions[session_id]
        
        # Add user input to conversation
        session.conversation_history.append({
            "role": "user",
            "content": user_input,
            "timestamp": datetime.now()
        })
        
        # Get AI response
        response = await self.llm_provider.get_best_response(
            f"Continue maintenance diagnosis conversation. "
            f"Machine: {session.machine_id}. "
            f"User says: {user_input}"
        )
        
        # Add AI response to conversation
        session.conversation_history.append({
            "role": "assistant",
            "content": response,
            "timestamp": datetime.now()
        })
        
        return response, session

class EnhancedVisionProcessor:
    def __init__(self, llm_provider: EnhancedLLMProvider):
        self.llm_provider = llm_provider
        self.storage_dir = Path("visual_uploads")
        self.storage_dir.mkdir(exist_ok=True)
    
    def save_visual_file(self, file_path: str, machine_id: str) -> str:
        """Save visual file and return new path"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    async def analyze_with_llama_vision(self, file_path: str, machine_context: str) -> str:
        """Analyze visual content with LLM"""
        prompt = f"""
        VISUAL INSPECTION ANALYSIS
        
        Machine Context: {machine_context}
        File: {file_path}
        
        Please analyze this visual inspection and identify:
        1. Visible defects or anomalies
        2. Wear patterns
        3. Potential safety concerns
        4. Maintenance recommendations
        
        Provide a detailed technical analysis.
        """
        
        return await self.llm_provider.get_best_response(prompt)

# Enhanced Smart Maintenance System
class EnhancedSmartMaintenanceSystem:
    def __init__(self):
        self.knowledge_base = EnhancedMaintenanceKnowledgeBase()
        self.llm_provider = EnhancedLLMProvider()
        self.diagnostic_engine = EnhancedDiagnosticEngine(self.knowledge_base, self.llm_provider)
        self.document_processor = DocumentProcessor()
        self.vision_processor = EnhancedVisionProcessor(self.llm_provider)
        self.current_session = None
    
    def display_menu(self):
        """Display interactive menu"""
        print("\n" + "="*60)
        print("SMART MAINTENANCE DIAGNOSTIC SYSTEM")
        print("="*60)
        print("1. Add New Equipment")
        print("2. Start Equipment Diagnosis")
        print("3. Upload Machine Documents")
        print("4. Upload Visual Inspection (Images/Videos)")
        print("5. Continue Active Diagnosis Session")
        print("6. View Equipment List")
        print("7. View Diagnosis History")
        print("8. Test LLM Connection")
        print("9. Change LLM Provider")
        print("10. Chat with AI Agent")
        print("11. Exit")
        print("-" * 60)
    
    def get_equipment_input(self) -> Dict:
        """Get equipment information from user"""
        print("\nEQUIPMENT REGISTRATION")
        print("-" * 40)
        
        equipment_id = input("Equipment ID: ").strip()
        if not equipment_id:
            print("Equipment ID is required!")
            return None
        
        name = input("Equipment Name: ").strip()
        equipment_type = input("Equipment Type (e.g., Pump, Motor, Compressor): ").strip()
        model = input("Model Number: ").strip()
        manufacturer = input("Manufacturer: ").strip()
        installation_date = input("Installation Date (YYYY-MM-DD) [Enter for today]: ").strip()
        last_maintenance = input("Last Maintenance Date (YYYY-MM-DD) [Enter for 'Never']: ").strip()
        
        if not installation_date:
            installation_date = datetime.now().strftime('%Y-%m-%d')
        if not last_maintenance:
            last_maintenance = 'Never'
        
        return {
            'id': equipment_id,
            'name': name or 'Unknown',
            'type': equipment_type or 'Unknown',
            'model': model or 'Unknown',
            'manufacturer': manufacturer or 'Unknown',
            'installation_date': installation_date,
            'last_maintenance': last_maintenance
        }
    
    def get_file_paths(self, file_type: str) -> List[str]:
        """Get file paths from user input"""
        files = []
        print(f"\nUPLOAD {file_type.upper()}")
        print("-" * 40)
        print("Enter file paths (one per line). Press Enter twice to finish:")
        
        while True:
            file_path = input(f"{file_type} file path: ").strip()
            if not file_path:
                break
            
            if os.path.exists(file_path):
                files.append(file_path)
                print(f"Added: {file_path}")
            else:
                print(f"File not found: {file_path}")
                retry = input("Try again? (y/n): ").lower().startswith('y')
                if retry:
                    continue
        
        return files
    
    async def add_equipment_interactive(self):
        """Interactive equipment addition"""
        equipment_data = self.get_equipment_input()
        if equipment_data:
            success = self.knowledge_base.add_equipment(equipment_data)
            if success:
                print(f"Equipment '{equipment_data['id']}' added successfully!")
            else:
                print("Failed to add equipment.")
        else:
            print("Equipment addition cancelled.")
    
    async def upload_documents_interactive(self):
        """Interactive document upload"""
        equipment_list = self.knowledge_base.get_all_equipment()
        if not equipment_list:
            print("No equipment found. Please add equipment first.")
            return
        
        print("\nAvailable Equipment:")
        for i, eq in enumerate(equipment_list, 1):
            print(f"{i}. {eq['id']} - {eq['name']}")
        
        try:
            choice = int(input("\nSelect equipment number: ")) - 1
            if 0 <= choice < len(equipment_list):
                machine_id = equipment_list[choice]['id']
                
                file_paths = self.get_file_paths("document")
                if file_paths:
                    uploaded_count = 0
                    for file_path in file_paths:
                        try:
                            # Process and save document
                            document = self.document_processor.process_document(file_path, machine_id)
                            # Save to storage
                            saved_path = self.document_processor.save_file(file_path, machine_id)
                            document.file_path = saved_path
                            
                            # Add to knowledge base
                            success = self.knowledge_base.add_machine_document(document)
                            if success:
                                uploaded_count += 1
                                print(f"Document processed: {file_path}")
                        except Exception as e:
                            print(f"Error processing {file_path}: {e}")
                    
                    print(f"{uploaded_count} documents uploaded successfully!")
                else:
                    print("No valid files provided.")
            else:
                print("Invalid selection.")
        except ValueError:
            print("Invalid input. Please enter a number.")
    
    async def upload_visual_inspection_interactive(self):
        """Interactive visual inspection upload"""
        equipment_list = self.knowledge_base.get_all_equipment()
        if not equipment_list:
            print("No equipment found. Please add equipment first.")
            return
        
        print("\nAvailable Equipment:")
        for i, eq in enumerate(equipment_list, 1):
            print(f"{i}. {eq['id']} - {eq['name']}")
    
        try:
            choice = int(input("\nSelect equipment number: ")) - 1
            if 0 <= choice < len(equipment_list):
                machine_id = equipment_list[choice]['id']
                
                file_paths = self.get_file_paths("visual (image/video)")
                if file_paths:
                    uploaded_count = 0
                    for file_path in file_paths:
                        try:
                            # Save visual file
                            saved_path = self.vision_processor.save_visual_file(file_path, machine_id)
                            
                            # Analyze visual content
                            print(f"Analyzing {file_path}...")
                            ai_analysis = await self.vision_processor.analyze_with_llama_vision(
                                saved_path, 
                                f"Machine: {equipment_list[choice]['name']} ({machine_id})"
                            )
                            
                            # Create visual inspection record
                            file_ext = Path(file_path).suffix.lower()
                            inspection_id = f"VIS_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                            
                            inspection = VisualInspection(
                                inspection_id=inspection_id,
                                machine_id=machine_id,
                                file_path=saved_path,
                                file_type="video" if file_ext in ['.mp4', '.avi', '.mov'] else "image",
                                defects_detected=[],
                                timestamp=datetime.now(),
                                ai_analysis=ai_analysis
                            )
                            
                            self.knowledge_base.add_visual_inspection(inspection)
                            uploaded_count += 1
                            print(f"Visual analysis completed for {file_path}")
                        except Exception as e:
                            print(f"Error processing {file_path}: {e}")
                    
                    print(f"{uploaded_count} visual files uploaded successfully!")
                else:
                    print("No valid files provided.")
            else:
                print("Invalid selection.")
        except ValueError:
            print("Invalid input. Please enter a number.")
    async def start_diagnosis_interactive(self):
        """Interactive diagnosis start"""
        equipment_list = self.knowledge_base.get_all_equipment()
        if not equipment_list:
            print("No equipment found. Please add equipment first.")
            return
        
        print("\nAvailable Equipment:")
        for i, eq in enumerate(equipment_list, 1):
            print(f"{i}. {eq['id']} - {eq['name']}")
        
        try:
            choice = int(input("\nSelect equipment number: ")) - 1
            if 0 <= choice < len(equipment_list):
                machine_id = equipment_list[choice]['id']
                equipment_name = equipment_list[choice]['name']
                
                print(f"\nSTARTING DIAGNOSIS FOR: {equipment_name}")
                print("-" * 50)
                
                # Step 1: Get problem description
                print("STEP 1: Problem Description")
                problem_description = input("Describe the problem you're experiencing: ").strip()
                if not problem_description:
                    print("Problem description is required.")
                    return
                
                # Step 2: Ask for visual uploads
                print("\nSTEP 2: Visual Inspection (Optional)")
                visual_upload = input("Do you have images or videos of the equipment? (y/n): ").lower().startswith('y')
                
                if visual_upload:
                    print("Please upload visual inspection files:")
                    file_paths = self.get_file_paths("visual (image/video)")
                    if file_paths:
                        for file_path in file_paths:
                            try:
                                saved_path = self.vision_processor.save_visual_file(file_path, machine_id)
                                print(f"Analyzing {file_path}...")
                                ai_analysis = await self.vision_processor.analyze_with_llama_vision(
                                    saved_path, f"Machine: {equipment_name} ({machine_id})"
                                )
                                
                                file_ext = Path(file_path).suffix.lower()
                                inspection_id = f"VIS_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                                
                                inspection = VisualInspection(
                                    inspection_id=inspection_id,
                                    machine_id=machine_id,
                                    file_path=saved_path,
                                    file_type="video" if file_ext in ['.mp4', '.avi', '.mov'] else "image",
                                    defects_detected=[],
                                    timestamp=datetime.now(),
                                    ai_analysis=ai_analysis
                                )
                                
                                self.knowledge_base.add_visual_inspection(inspection)
                                print(f"Visual analysis completed for {file_path}")
                            except Exception as e:
                                print(f"Error processing {file_path}: {e}")
                
                # Step 3: Ask for document uploads
                print("\nSTEP 3: Machine Documentation (Optional)")
                doc_upload = input("Do you have manuals, maintenance records, or other documents? (y/n): ").lower().startswith('y')
                
                if doc_upload:
                    print("Please upload machine documentation:")
                    file_paths = self.get_file_paths("document")
                    if file_paths:
                        for file_path in file_paths:
                            try:
                                document = self.document_processor.process_document(file_path, machine_id)
                                saved_path = self.document_processor.save_file(file_path, machine_id)
                                document.file_path = saved_path
                                
                                self.knowledge_base.add_machine_document(document)
                                print(f"Document processed: {file_path}")
                            except Exception as e:
                                print(f"Error processing {file_path}: {e}")
                
                # Step 4: Analyze all data and create problem statement
                print("\nSTEP 4: Analyzing Data and Generating Problem Statement")
                print("Please wait while AI analyzes all available information...")
                
                # Get all available data
                machine_docs = self.knowledge_base.get_machine_documents(machine_id)
                visual_data = self.knowledge_base.get_visual_inspections(machine_id)
                
                # Generate comprehensive problem statement
                problem_statement_prompt = f"""
                MACHINE ANALYSIS AND PROBLEM STATEMENT GENERATION
                
                Equipment: {equipment_name} (ID: {machine_id})
                Equipment Type: {equipment_list[choice]['type']}
                User Problem Description: {problem_description}
                
                Available Documentation: {len(machine_docs)} documents
                Visual Inspections: {len(visual_data)} files
                
                Based on the user's problem description and all available data, please:
                1. Create a detailed technical problem statement
                2. Identify key symptoms and their significance
                3. List relevant equipment specifications and operating conditions
                4. Summarize visual inspection findings (if any)
                5. Reference relevant documentation sections (if any)
                6. Prepare the problem for comprehensive diagnosis
                
                Format the output as a structured problem statement that can be used for detailed diagnosis.
                """
                
                problem_statement = await self.llm_provider.analyze_with_context(
                    problem_statement_prompt, machine_docs, visual_data
                )
                
                print("\nGENERATED PROBLEM STATEMENT:")
                print("=" * 50)
                print(problem_statement)
                print("=" * 50)
                
                # Step 5: Perform comprehensive diagnosis
                print("\nSTEP 5: Comprehensive Diagnosis")
                print("Performing detailed analysis...")
                
                # Start diagnosis session
                session = await self.diagnostic_engine.start_interactive_diagnosis(machine_id)
                self.current_session = session
                
                # Perform comprehensive diagnosis with enhanced problem statement
                diagnosis = await self.diagnostic_engine.comprehensive_diagnosis(
                    machine_id=machine_id,
                    user_description=f"Original Problem: {problem_description}\n\nDetailed Problem Statement: {problem_statement}"
                )
                
                # Display results
                self.display_diagnosis_results(diagnosis)
                
                # Step 6: Interactive follow-up
                print("\nSTEP 6: Interactive Follow-up")
                await self.continue_interactive_diagnosis(session.session_id)
                
            else:
                print("Invalid selection.")
        except ValueError:
            print("Invalid input. Please enter a number.")
    
    async def continue_interactive_diagnosis(self, session_id: str, initial_problem: str = ""):
        """Continue interactive diagnosis conversation"""
        print("\nINTERACTIVE DIAGNOSIS SESSION")
        print("Type 'exit' to end session, 'upload' to add more files")
        print("-" * 50)
        
        while True:
            user_input = input("\nYour input: ").strip()
            
            if user_input.lower() == 'exit':
                print("Diagnosis session ended.")
                break
            elif user_input.lower() == 'upload':
                print("\nChoose file type to upload:")
                print("1. Documents")
                print("2. Visual Inspection (Images/Videos)")
                choice = input("Choice (1-2): ").strip()
                
                if choice == "1":
                    await self.upload_documents_interactive()
                elif choice == "2":
                    await self.upload_visual_inspection_interactive()
                continue
            elif not user_input:
                continue
            
            try:
                print("Analyzing...")
                response, updated_session = await self.diagnostic_engine.continue_diagnosis(
                    session_id, user_input
                )
                
                print(f"\nAI Response:")
                print("-" * 30)
                print(response)
                
            except Exception as e:
                print(f"Error: {e}")
    
    def display_diagnosis_results(self, diagnosis: FaultDiagnosis):
        """Display diagnosis results in formatted way"""
        print("\n" + "="*60)
        print("DIAGNOSIS RESULTS")
        print("="*60)
        print(f"Fault ID: {diagnosis.fault_id}")
        print(f"Machine: {diagnosis.machine_id}")
        print(f"Severity: {diagnosis.severity.upper()}")
        print(f"Confidence: {diagnosis.confidence:.1%}")
        print(f"\nRoot Cause:")
        print(f"  {diagnosis.root_cause}")
        
        if diagnosis.recommended_actions:
            print(f"\nRecommended Actions:")
            for i, action in enumerate(diagnosis.recommended_actions, 1):
                print(f"  {i}. {action}")
        
        print(f"\nEstimated Downtime: {diagnosis.estimated_downtime} hours")
        
        if diagnosis.supporting_evidence:
            print(f"\nSupporting Evidence:")
            for evidence in diagnosis.supporting_evidence:
                print(f"  - {evidence}")
        
        print("="*60)
    
    def view_equipment_list(self):
        """Display all equipment"""
        equipment_list = self.knowledge_base.get_all_equipment()
        if not equipment_list:
            print("No equipment found.")
            return
        
        print("\nEQUIPMENT LIST")
        print("-" * 80)
        print(f"{'ID':<15} {'Name':<25} {'Type':<20} {'Last Maintenance':<15}")
        print("-" * 80)
        
        for eq in equipment_list:
            print(f"{eq['id']:<15} {eq['name']:<25} {eq['type']:<20} {eq['last_maintenance']:<15}")
    
    def view_diagnosis_history(self):
        """Display diagnosis history"""
        print("Diagnosis history feature coming soon...")
    
    def change_llm_provider(self):
        """Change LLM provider"""
        print(f"\nCurrent LLM Provider: {self.llm_provider.get_current_provider()}")
        print("\nAvailable providers:")
        print("1. auto (tries all available providers)")
        print("2. groq")
        print("3. gemini")
        print("4. huggingface")
        print("5. local (offline fallback)")
        
        choice = input("\nSelect provider (1-5): ").strip()
        
        provider_map = {
            "1": "auto",
            "2": "groq",
            "3": "gemini",
            "4": "huggingface",
            "5": "local"
        }
        
        if choice in provider_map:
            self.llm_provider.set_provider(provider_map[choice])
        else:
            print("Invalid selection.")
    
    async def chat_with_ai(self):
        """Direct chat with AI agent"""
        print("\nCHAT WITH AI MAINTENANCE AGENT")
        print("Type 'exit' to end chat session")
        print("-" * 40)
        print(f"Current LLM Provider: {self.llm_provider.get_current_provider()}")
        print("-" * 40)
        
        while True:
            user_input = input("\nYou: ").strip()
            
            if user_input.lower() == 'exit':
                print("Chat session ended.")
                break
            elif not user_input:
                continue
            
            try:
                print("AI Agent: ", end="")
                response = await self.llm_provider.get_best_response(
                    f"You are an expert industrial maintenance engineer. "
                    f"Provide helpful, accurate, and detailed responses about maintenance, "
                    f"troubleshooting, and equipment issues. "
                    f"User question: {user_input}"
                )
                print(response)
                
            except Exception as e:
                print(f"Error: {e}")
    
    async def run_interactive_system(self):
        """Main interactive system loop"""
        print("Starting Smart Maintenance Diagnostic System...")
    async def test_connection(self):
        """Test LLM connection"""
        return await self.llm_provider.get_best_response("Test connection. Respond with 'Connection successful'")
    
    async def run_interactive_system(self):
        """Main interactive system loop"""
        print("Starting Smart Maintenance Diagnostic System...")
        
        # Test connection on startup
        try:
            test_response = await self.test_connection()
            print("LLM Connection successful!")
        except Exception as e:
            print(f"LLM Connection warning: {e}")
        
        while True:
            try:
                self.display_menu()
                choice = input("Select option (1-11): ").strip()
                
                if choice == "1":
                    await self.add_equipment_interactive()
                elif choice == "2":
                    await self.start_diagnosis_interactive()
                elif choice == "3":
                    await self.upload_documents_interactive()
                elif choice == "4":
                    await self.upload_visual_inspection_interactive()
                elif choice == "5":
                    if self.current_session:
                        await self.continue_interactive_diagnosis(self.current_session.session_id)
                    else:
                        print("No active diagnosis session.")
                elif choice == "6":
                    self.view_equipment_list()
                elif choice == "7":
                    self.view_diagnosis_history()
                elif choice == "8":
                    print("Testing LLM connection...")
                    test_response = await self.test_connection()
                    print(f"Response: {test_response[:300]}...")
                elif choice == "9":
                    self.change_llm_provider()
                elif choice == "10":
                    await self.chat_with_ai()
                elif choice == "11":
                    print("Goodbye!")
                    break
                else:
                    print("Invalid option. Please select 1-11.")
                
                input("\nPress Enter to continue...")
                
            except KeyboardInterrupt:
                print("\n\nSystem interrupted. Goodbye!")
                break
            except Exception as e:
                print(f"Unexpected error: {e}")
                input("Press Enter to continue...")
    
    async def close(self):
        """Cleanup system resources"""
        await self.llm_provider.close()

# Interactive main function
async def main():
    """Interactive Smart Maintenance System"""
    system = EnhancedSmartMaintenanceSystem()
    
    try:
        await system.run_interactive_system()
    except Exception as e:
        print(f"System error: {e}")
    finally:
        await system.close()

if __name__ == "__main__":
    asyncio.run(main())