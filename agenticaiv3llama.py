import asyncio
import json
import logging
import sqlite3
import os
import cv2
import numpy as np
import re
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from enum import Enum

import aiohttp
from dotenv import load_dotenv
import PyPDF2
import docx
from pathlib import Path
import hashlib
import shutil
load_dotenv()

class LLMConfig:
    # """Configuration for LLM providers"""
    # LLAMA_API_KEY = os.getenv("LLAMA_API_KEY", "your-llama-api-key")
    # LLAMA_BASE_URL = os.getenv("LLAMA_BASE_URL", "https://api.openai.com/v1")
    
    # Alternative free providers
    GROQ_API_KEY = os.getenv("GROQ_API_KEY", "your-groq-api-key")
    HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY", "your-hf-api-key")
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "your-gemini-api-key")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enhanced Data Models
@dataclass
class SensorData:
    sensor_id: str
    timestamp: datetime
    temperature: float
    vibration: float
    pressure: float
    humidity: float
    sound_level: float

@dataclass
class MachineDocument:
    doc_id: str
    machine_id: str
    file_path: str
    file_type: str
    content: str
    upload_date: datetime
    summary: str = ""

@dataclass
class VisualInspection:
    inspection_id: str
    machine_id: str
    file_path: str
    file_type: str  # image or video
    defects_detected: List[Dict]
    timestamp: datetime
    ai_analysis: str = ""

@dataclass
class FaultDiagnosis:
    fault_id: str
    machine_id: str
    severity: str
    confidence: float
    description: str
    root_cause: str
    recommended_actions: List[str]
    estimated_downtime: int
    supporting_evidence: List[str]  # References to documents/images
    conversation_history: Optional[List[Dict]] = None

@dataclass
class DiagnosisSession:
    session_id: str
    machine_id: str
    start_time: datetime
    current_diagnosis: Optional[FaultDiagnosis]
    conversation_history: List[Dict]
    uploaded_files: List[str]
    status: str  # active, resolved, escalated

class MaintenanceType(Enum):
    PREVENTIVE = "preventive"
    CORRECTIVE = "corrective"
    PREDICTIVE = "predictive"
    EMERGENCY = "emergency"

@dataclass
class MaintenanceTask:
    task_id: str
    equipment_id: str
    task_type: MaintenanceType
    priority: int
    description: str
    scheduled_date: datetime
    estimated_duration: int
    required_parts: List[str]
    assigned_technician: Optional[str] = None

# Enhanced Document Processing
class DocumentProcessor:
    def __init__(self, storage_dir: str = "uploads"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
    def save_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded file and return new path"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def process_document(self, file_path: str, machine_id: str) -> MachineDocument:
        """Process uploaded document and extract content"""
        file_path_obj = Path(file_path)
        file_type = file_path_obj.suffix.lower()
        
        # Extract text based on file type
        if file_type == '.pdf':
            content = self.extract_text_from_pdf(file_path)
        elif file_type == '.docx':
            content = self.extract_text_from_docx(file_path)
        elif file_type in ['.txt', '.md']:
            content = self.extract_text_from_txt(file_path)
        else:
            content = f"Unsupported file type: {file_type}"
        
        # Generate document ID
        doc_id = hashlib.md5(f"{machine_id}_{file_path_obj.name}_{datetime.now()}".encode()).hexdigest()[:12]
        
        return MachineDocument(
            doc_id=doc_id,
            machine_id=machine_id,
            file_path=file_path,
            file_type=file_type,
            content=content,
            upload_date=datetime.now()
        )

# Enhanced Visual Processing
# class EnhancedVisionProcessor:
#     def __init__(self, llm_provider, storage_dir: str = "visual_uploads"):
#         self.llm_provider = llm_provider
#         self.storage_dir = Path(storage_dir)
#         self.storage_dir.mkdir(exist_ok=True)
    
#     def save_visual_file(self, file_path: str, machine_id: str) -> str:
#         """Save uploaded visual file"""
#         original_path = Path(file_path)
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
#         new_path = self.storage_dir / new_filename
        
#         import shutil
#         shutil.copy2(file_path, new_path)
#         return str(new_path)
    
#     def analyze_image(self, image_path: str) -> Dict[str, Any]:
#         """Enhanced image analysis with OpenCV"""
#         try:
#             image = cv2.imread(image_path)
#             if image is None:
#                 return {"error": "Could not load image"}
            
#             # Multiple analysis techniques
#             analysis_results = {}
            
#             # 1. Basic defect detection
#             gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
#             edges = cv2.Canny(gray, 50, 150)
#             contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
#             defects = []
#             for contour in contours:
#                 area = cv2.contourArea(contour)
#                 if area > 500:  # Threshold for significant defects
#                     x, y, w, h = cv2.boundingRect(contour)
#                     defects.append({
#                         'location': {'x': int(x), 'y': int(y), 'width': int(w), 'height': int(h)},
#                         'area': float(area),
#                         'type': 'potential_defect'
#                     })
            
#             # 2. Corrosion detection (rust-like colors)
#             hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
#             rust_lower = np.array([0, 50, 50])
#             rust_upper = np.array([20, 255, 255])
#             rust_mask = cv2.inRange(hsv, rust_lower, rust_upper)
#             rust_area = cv2.countNonZero(rust_mask)
            
#             # 3. Crack detection using morphological operations
#             kernel = np.ones((3,3), np.uint8)
#             morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
#             crack_edges = cv2.Canny(morph, 100, 200)
#             crack_lines = cv2.HoughLinesP(crack_edges, 1, np.pi/180, threshold=50, minLineLength=30, maxLineGap=10)
            
#             analysis_results = {
#                 'defects_found': len(defects),
#                 'defects': defects,
#                 'rust_area_pixels': int(rust_area),
#                 'potential_cracks': len(crack_lines) if crack_lines is not None else 0,
#                 'image_dimensions': {'width': image.shape[1], 'height': image.shape[0]},
#                 'analysis_timestamp': datetime.now().isoformat()
#             }
            
#             return analysis_results
            
#         except Exception as e:
#             logger.error(f"Vision processing error: {e}")
#             return {"error": str(e)}
    
#     def analyze_video(self, video_path: str) -> Dict[str, Any]:
#         """Basic video analysis for defect detection"""
#         try:
#             cap = cv2.VideoCapture(video_path)
#             if not cap.isOpened():
#                 return {"error": "Could not open video"}
            
#             frame_count = 0
#             total_defects = 0
#             frame_analyses = []
            
#             # Analyze every 30th frame to avoid processing every frame
#             while True:
#                 ret, frame = cap.read()
#                 if not ret:
#                     break
                
#                 if frame_count % 30 == 0:  # Process every 30th frame
#                     # Save frame temporarily
#                     temp_frame_path = f"temp_frame_{frame_count}.jpg"
#                     cv2.imwrite(temp_frame_path, frame)
                    
#                     # Analyze frame
#                     frame_analysis = self.analyze_image(temp_frame_path)
#                     if 'defects_found' in frame_analysis:
#                         total_defects += frame_analysis['defects_found']
#                         frame_analyses.append({
#                             'frame_number': frame_count,
#                             'defects': frame_analysis['defects_found'],
#                             'timestamp': frame_count / cap.get(cv2.CAP_PROP_FPS) if cap.get(cv2.CAP_PROP_FPS) > 0 else 0
#                         })
                    
#                     # Clean up temp file
#                     if os.path.exists(temp_frame_path):
#                         os.remove(temp_frame_path)
                
#                 frame_count += 1
            
#             cap.release()
            
#             return {
#                 'total_frames': frame_count,
#                 'frames_analyzed': len(frame_analyses),
#                 'total_defects_detected': total_defects,
#                 'frame_analyses': frame_analyses,
#                 'analysis_timestamp': datetime.now().isoformat()
#             }
            
#         except Exception as e:
#             logger.error(f"Video processing error: {e}")
#             return {"error": str(e)}
    
#     async def analyze_with_llama_vision(self, file_path: str, machine_context: str = "") -> str:
#         """Analyze visual content with LLaMA (text-based analysis of CV results)"""
#         try:
#             # Get computer vision analysis first
#             file_ext = Path(file_path).suffix.lower()
#             is_video = file_ext in ['.mp4', '.avi', '.mov', '.mkv']
            
#             if is_video:
#                 cv_analysis = self.analyze_video(file_path)
#             else:
#                 cv_analysis = self.analyze_image(file_path)
            
#             prompt = f"""
#             As an expert industrial maintenance engineer, analyze the following computer vision inspection results:
            
#             Machine Context: {machine_context}
#             File Type: {"Video" if is_video else "Image"}
#             File Path: {file_path}
            
#             COMPUTER VISION ANALYSIS RESULTS:
#             {json.dumps(cv_analysis, indent=2)}
            
#             Based on these computer vision findings, please provide:
#             1. Interpretation of detected defects and their significance
#             2. Assessment of rust/corrosion areas (if any)
#             3. Evaluation of potential cracks or structural issues
#             4. Overall condition assessment
#             5. Immediate safety concerns (if any)
#             6. Recommended maintenance actions
#             7. Priority level for addressing identified issues
            
#             Provide specific, actionable recommendations based on the computer vision analysis results.
#             Consider the machine context and typical failure modes for this type of equipment.
#             """
            
#             response = await self.llm_provider.get_best_response(prompt)
#             return response
            
#         except Exception as e:
#             logger.error(f"LLaMA vision analysis error: {e}")
#             return f"Visual analysis error: {e}"



# Enhanced LLM Provider with multiple providers
# class EnhancedLLMProvider:
#     def __init__(self):
#         self.init_llama()
    
#     def init_llama(self):
#         """Initialize LLaMA provider via OpenAI"""
#         try:
#             if LLMConfig.LLAMA_API_KEY and LLMConfig.LLAMA_API_KEY != "your-llama-api-key":
#                 self.llama_configured = True
#                 self.api_key = LLMConfig.LLAMA_API_KEY
#                 self.base_url = LLMConfig.LLAMA_BASE_URL or "https://api.openai.com/v1"
#                 # Use OpenAI's model names for LLaMA access
#                 self.model_name = "gpt-4o"  # or "gpt-3.5-turbo" for cost efficiency
#                 logger.info("LLaMA model initialized via OpenAI successfully")
#             else:
#                 self.llama_configured = False
#                 logger.warning("LLaMA/OpenAI API key not configured")
#         except Exception as e:
#             logger.warning(f"Error initializing LLaMA via OpenAI: {e}")
#             self.llama_configured = False
    
#     async def call_llama(self, prompt: str) -> str:
#         """Call LLaMA API"""
#         try:
#             if not self.llama_configured:
#                 return "LLaMA API not configured. Please set LLAMA_API_KEY environment variable."
            
#             headers = {
#                 "Authorization": f"Bearer {self.api_key}",
#                 "Content-Type": "application/json"
#             }
            
#             payload = {
#                 "model": self.model_name,
#                 "messages": [
#                     {
#                         "role": "user",
#                         "content": prompt
#                     }
#                 ],
#                 "max_tokens": 2048,
#                 "temperature": 0.7,
#                 "top_p": 0.9
#             }
            
#             async with aiohttp.ClientSession() as session:
#                 async with session.post(
#                     f"{self.base_url}/chat/completions",
#                     headers=headers,
#                     json=payload,
#                     timeout=aiohttp.ClientTimeout(total=60)
#                 ) as response:
#                     if response.status == 200:
#                         data = await response.json()
#                         return data["choices"][0]["message"]["content"].strip()
#                     else:
#                         error_text = await response.text()
#                         logger.error(f"LLaMA API error {response.status}: {error_text}")
#                         return f"LLaMA API error: {response.status} - {error_text}"
                        
#         except Exception as e:
#             logger.error(f"LLaMA API error: {e}")
#             return f"Error calling LLaMA: {e}"
    
#     async def call_gemini(self, prompt: str) -> str:
#         """Compatibility method - redirects to LLaMA"""
#         return await self.call_llama(prompt)
    
#     async def get_best_response(self, prompt: str) -> str:
#         """Get response from LLaMA"""
#         return await self.call_llama(prompt)
    
#     async def analyze_with_context(self, prompt: str, context_docs: List[MachineDocument], visual_data: Optional[List[VisualInspection]] = None) -> str:
#         """Analyze with machine-specific context"""
#         enhanced_prompt = f"""
#         MACHINE-SPECIFIC CONTEXT:
#         {self._format_document_context(context_docs)}
        
#         VISUAL INSPECTION DATA:
#         {self._format_visual_context(visual_data) if visual_data else "No visual data available"}
        
#         MAINTENANCE EXPERT ANALYSIS REQUEST:
#         {prompt}
        
#         Please provide analysis based on the specific machine documentation and visual inspection data provided above.
#         Reference specific sections of the documentation when relevant.
#         """
        
#         return await self.get_best_response(enhanced_prompt)
    
#     def _format_document_context(self, docs: List[MachineDocument]) -> str:
#         """Format document context for LLM"""
#         if not docs:
#             return "No machine documentation available"
        
#         context = ""
#         for doc in docs:
#             context += f"""
#             Document: {doc.file_path}
#             Type: {doc.file_type}
#             Content Summary: {doc.content[:1000]}...
#             ---
#             """
#         return context
    
#     def _format_visual_context(self, visual_data: List[VisualInspection]) -> str:
#         """Format visual inspection context"""
#         if not visual_data:
#             return "No visual inspection data available"
        
#         context = ""
#         for inspection in visual_data:
#             context += f"""
#             Visual Inspection: {inspection.file_path}
#             Type: {inspection.file_type}
#             Defects Detected: {len(inspection.defects_detected)}
#             AI Analysis: {inspection.ai_analysis[:500]}...
#             ---
#             """
#         return context
    
#     async def close(self):
#         """Cleanup resources"""
#         pass
class EnhancedLLMProvider:
    def __init__(self):
        self.init_providers()
    
    def init_providers(self):
        """Initialize multiple LLM providers"""
        # # OpenAI/LLaMA
        # self.openai_configured = (LLMConfig.LLAMA_API_KEY and 
        #                         LLMConfig.LLAMA_API_KEY != "your-llama-api-key")
        
        # Groq (Free tier available)
        self.groq_configured = (LLMConfig.GROQ_API_KEY and 
                              LLMConfig.GROQ_API_KEY != "your-groq-api-key")
        
        # Hugging Face (Free tier available)
        self.hf_configured = (LLMConfig.HUGGINGFACE_API_KEY and 
                            LLMConfig.HUGGINGFACE_API_KEY != "your-hf-api-key")
        
        # Gemini (Free tier available)
        self.gemini_configured = (LLMConfig.GEMINI_API_KEY and 
                                LLMConfig.GEMINI_API_KEY != "your-gemini-api-key")
        
        logger.info(f"Providers configured: "
                   f"Groq={self.groq_configured}, HF={self.hf_configured}, "
                   f"Gemini={self.gemini_configured}")
    
    async def call_groq(self, prompt: str) -> str:
        """Call Groq API (free tier available)"""
        try:
            if not self.groq_configured:
                return "Groq API not configured"
            
            headers = {
                "Authorization": f"Bearer {LLMConfig.GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": "llama3-8b-8192",  # Free model
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 2048,
                "temperature": 0.7
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        return data["choices"][0]["message"]["content"].strip()
                    else:
                        error_text = await response.text()
                        logger.error(f"Groq API error {response.status}: {error_text}")
                        return f"Groq API error: {response.status}"
                        
        except Exception as e:
            logger.error(f"Groq API error: {e}")
            return f"Error calling Groq: {e}"
    
    async def call_huggingface(self, prompt: str) -> str:
        """Call Hugging Face API (free tier available)"""
        try:
            if not self.hf_configured:
                return "Hugging Face API not configured"
            
            headers = {
                "Authorization": f"Bearer {LLMConfig.HUGGINGFACE_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": 1000,
                    "temperature": 0.7,
                    "return_full_text": False
                }
            }
            
            # Using a free model
            model_url = "https://api-inference.huggingface.co/models/microsoft/DialoGPT-large"
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    model_url,
                    headers=headers,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status == 200:
                        data = await response.json()
                        if isinstance(data, list) and len(data) > 0:
                            return data[0].get("generated_text", "No response generated")
                        return str(data)
                    else:
                        error_text = await response.text()
                        logger.error(f"HF API error {response.status}: {error_text}")
                        return f"HF API error: {response.status}"
                        
        except Exception as e:
            logger.error(f"HF API error: {e}")
            return f"Error calling Hugging Face: {e}"
        
        async def call_gemini(self, prompt: str) -> str:
            """Call Gemini API (free tier available)"""
            try:
                if not self.gemini_configured:
                    return "Gemini API not configured"
                
                url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={LLMConfig.GEMINI_API_KEY}"
                
                payload = {
                    "contents": [{
                        "parts": [{"text": prompt}]
                    }],
                    "generationConfig": {
                        "temperature": 0.7,
                        "maxOutputTokens": 2048
                    }
                }
                
                async with aiohttp.ClientSession() as session:
                    async with session.post(
                        url,
                        json=payload,
                        timeout=aiohttp.ClientTimeout(total=60)
                    ) as response:
                        if response.status == 200:
                            data = await response.json()
                            if 'candidates' in data and len(data['candidates']) > 0:
                                content = data['candidates'][0]['content']['parts'][0]['text']
                                return content.strip()
                            return "No response generated"
                        else:
                            error_text = await response.text()
                            logger.error(f"Gemini API error {response.status}: {error_text}")
                            return f"Gemini API error: {response.status}"
                            
            except Exception as e:
                logger.error(f"Gemini API error: {e}")
                return f"Error calling Gemini: {e}"
            
        # async def call_llama(self, prompt: str) -> str:
        #     """Call OpenAI API (original method)"""
        #     try:
        #         if not self.openai_configured:
        #             return "OpenAI API not configured"
                
        #         headers = {
        #             "Authorization": f"Bearer {LLMConfig.LLAMA_API_KEY}",
        #             "Content-Type": "application/json"
        #         }
                
        #         payload = {
        #             "model": "gpt-4o-mini",  # Use cheaper model
        #             "messages": [{"role": "user", "content": prompt}],
        #             "max_tokens": 1000,  # Reduce token usage
        #             "temperature": 0.7
        #         }
                
        #         async with aiohttp.ClientSession() as session:
        #             async with session.post(
        #                 f"{LLMConfig.LLAMA_BASE_URL}/chat/completions",
        #                 headers=headers,
        #                 json=payload,
        #                 timeout=aiohttp.ClientTimeout(total=60)
        #             ) as response:
        #                 if response.status == 200:
        #                     data = await response.json()
        #                     return data["choices"][0]["message"]["content"].strip()
        #                 else:
        #                     error_text = await response.text()
        #                     logger.error(f"OpenAI API error {response.status}: {error_text}")
        #                     return f"OpenAI API error: {response.status} - {error_text}"
                            
        #     except Exception as e:
        #         logger.error(f"OpenAI API error: {e}")
        #         return f"Error calling OpenAI: {e}"
    
    
    async def call_local_fallback(self, prompt: str) -> str:
        """Local fallback with pre-defined responses"""
        maintenance_keywords = {
            'bearing': "Common bearing failure causes: inadequate lubrication, misalignment, contamination, overloading, and fatigue. Recommended actions: check lubrication, inspect alignment, replace if worn.",
            'pump': "Common pump issues: cavitation, impeller wear, seal failure, misalignment. Check suction pressure, inspect impeller, verify shaft alignment.",
            'motor': "Common motor problems: overheating, bearing failure, winding damage. Check temperature, vibration levels, electrical connections.",
            'vibration': "Excessive vibration causes: imbalance, misalignment, bearing wear, loose components. Perform vibration analysis, check alignment, inspect bearings.",
            'noise': "Unusual noise sources: bearing wear, cavitation, loose parts, misalignment. Identify noise type and location for proper diagnosis.",
            'leak': "Common leak sources: seal failure, gasket degradation, pipe joints, valve packing. Inspect seals, replace gaskets, check torque specifications."
        }
        
        prompt_lower = prompt.lower()
        for keyword, response in maintenance_keywords.items():
            if keyword in prompt_lower:
                return f"MAINTENANCE ANALYSIS (Local Fallback): {response}"
        
        return "MAINTENANCE ANALYSIS (Local Fallback): Unable to analyze without specific equipment information. Please provide more details about the equipment type, symptoms, and operating conditions."
    
    async def get_best_response(self, prompt: str) -> str:
        """Try providers in order of preference"""
        providers = [
            # ("OpenAI", self.call_llama) if self.openai_configured else None,
            ("Groq", self.call_groq) if self.groq_configured else None,
            ("HuggingFace", self.call_huggingface) if self.hf_configured else None,
            ("Gemini", self.call_gemini) if self.gemini_configured else None,
            ("Local Fallback", self.call_local_fallback)
        ]
        
        # Filter out None providers
        providers = [p for p in providers if p is not None]
        
        for provider_name, provider_func in providers:
            try:
                logger.info(f"Trying {provider_name} provider...")
                response = await provider_func(prompt)
                
                # Check if response indicates an error
                if not any(error_term in response.lower() for error_term in 
                          ['error', 'quota', 'exceeded', 'not configured']):
                    logger.info(f"Success with {provider_name}")
                    return response
                else:
                    logger.warning(f"{provider_name} failed: {response[:100]}...")
                    
            except Exception as e:
                logger.error(f"{provider_name} provider failed: {e}")
                continue
        
        # If all providers fail, return local fallback
        return await self.call_local_fallback(prompt)
    
    
    async def analyze_with_context(self, prompt: str, context_docs: List[MachineDocument], visual_data: Optional[List[VisualInspection]] = None) -> str:
        """Analyze with machine-specific context"""
        enhanced_prompt = f"""
        MACHINE-SPECIFIC CONTEXT:
        {self._format_document_context(context_docs)}
        
        VISUAL INSPECTION DATA:
        {self._format_visual_context(visual_data) if visual_data else "No visual data available"}
        
        MAINTENANCE EXPERT ANALYSIS REQUEST:
        {prompt}
        
        Please provide analysis based on the specific machine documentation and visual inspection data provided above.
        Reference specific sections of the documentation when relevant.
        """
        
        return await self.get_best_response(enhanced_prompt)
    
    def _format_document_context(self, docs: List[MachineDocument]) -> str:
        """Format document context for LLM"""
        if not docs:
            return "No machine documentation available"
        
        context = ""
        for doc in docs:
            context += f"""
            Document: {doc.file_path}
            Type: {doc.file_type}
            Content Summary: {doc.content[:500]}...
            ---
            """
        return context
    
    def _format_visual_context(self, visual_data: List[VisualInspection]) -> str:
        """Format visual inspection context"""
        if not visual_data:
            return "No visual inspection data available"
        
        context = ""
        for inspection in visual_data:
            context += f"""
            Visual Inspection: {inspection.file_path}
            Type: {inspection.file_type}
            Defects Detected: {len(inspection.defects_detected)}
            AI Analysis: {inspection.ai_analysis[:300]}...
            ---
            """
        return context
    
    async def close(self):
        """Cleanup resources"""
        pass

class EnhancedVisionProcessor:
    def __init__(self, llm_provider: EnhancedLLMProvider, storage_dir: str = "visual_uploads"):
        self.llm_provider = llm_provider
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
    
    def save_visual_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded visual file"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        
        import shutil
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def analyze_image(self, image_path: str) -> Dict[str, Any]:
        """Enhanced image analysis with OpenCV"""
        try:
            image = cv2.imread(image_path)
            if image is None:
                return {"error": "Could not load image"}
            
            # Multiple analysis techniques
            analysis_results = {}
            
            # 1. Basic defect detection
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            edges = cv2.Canny(gray, 50, 150)
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            defects = []
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > 500:  # Threshold for significant defects
                    x, y, w, h = cv2.boundingRect(contour)
                    defects.append({
                        'location': {'x': int(x), 'y': int(y), 'width': int(w), 'height': int(h)},
                        'area': float(area),
                        'type': 'potential_defect'
                    })
            
            # 2. Corrosion detection (rust-like colors)
            hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
            rust_lower = np.array([0, 50, 50])
            rust_upper = np.array([20, 255, 255])
            rust_mask = cv2.inRange(hsv, rust_lower, rust_upper)
            rust_area = cv2.countNonZero(rust_mask)
            
            # 3. Crack detection using morphological operations
            kernel = np.ones((3,3), np.uint8)
            morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
            crack_edges = cv2.Canny(morph, 100, 200)
            crack_lines = cv2.HoughLinesP(crack_edges, 1, np.pi/180, threshold=50, minLineLength=30, maxLineGap=10)
            
            analysis_results = {
                'defects_found': len(defects),
                'defects': defects,
                'rust_area_pixels': int(rust_area),
                'potential_cracks': len(crack_lines) if crack_lines is not None else 0,
                'image_dimensions': {'width': image.shape[1], 'height': image.shape[0]},
                'analysis_timestamp': datetime.now().isoformat()
            }
            
            return analysis_results
            
        except Exception as e:
            logger.error(f"Vision processing error: {e}")
            return {"error": str(e)}
    
    def analyze_video(self, video_path: str) -> Dict[str, Any]:
        """Basic video analysis for defect detection"""
        try:
            cap = cv2.VideoCapture(video_path)
            if not cap.isOpened():
                return {"error": "Could not open video"}
            
            frame_count = 0
            total_defects = 0
            frame_analyses = []
            
            # Analyze every 30th frame to avoid processing every frame
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                
                if frame_count % 30 == 0:  # Process every 30th frame
                    # Save frame temporarily
                    temp_frame_path = f"temp_frame_{frame_count}.jpg"
                    cv2.imwrite(temp_frame_path, frame)
                    
                    # Analyze frame
                    frame_analysis = self.analyze_image(temp_frame_path)
                    if 'defects_found' in frame_analysis:
                        total_defects += frame_analysis['defects_found']
                        frame_analyses.append({
                            'frame_number': frame_count,
                            'defects': frame_analysis['defects_found'],
                            'timestamp': frame_count / cap.get(cv2.CAP_PROP_FPS) if cap.get(cv2.CAP_PROP_FPS) > 0 else 0
                        })
                    
                    # Clean up temp file
                    if os.path.exists(temp_frame_path):
                        os.remove(temp_frame_path)
                
                frame_count += 1
            
            cap.release()
            
            return {
                'total_frames': frame_count,
                'frames_analyzed': len(frame_analyses),
                'total_defects_detected': total_defects,
                'frame_analyses': frame_analyses,
                'analysis_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Video processing error: {e}")
            return {"error": str(e)}
    
    async def analyze_with_llama_vision(self, file_path: str, machine_context: str = "") -> str:
        """Analyze visual content with LLaMA (text-based analysis of CV results)"""
        try:
            # Get computer vision analysis first
            file_ext = Path(file_path).suffix.lower()
            is_video = file_ext in ['.mp4', '.avi', '.mov', '.mkv']
            
            if is_video:
                cv_analysis = self.analyze_video(file_path)
            else:
                cv_analysis = self.analyze_image(file_path)
            
            prompt = f"""
            As an expert industrial maintenance engineer, analyze the following computer vision inspection results:
            
            Machine Context: {machine_context}
            File Type: {"Video" if is_video else "Image"}
            File Path: {file_path}
            
            COMPUTER VISION ANALYSIS RESULTS:
            {json.dumps(cv_analysis, indent=2)}
            
            Based on these computer vision findings, please provide:
            1. Interpretation of detected defects and their significance
            2. Assessment of rust/corrosion areas (if any)
            3. Evaluation of potential cracks or structural issues
            4. Overall condition assessment
            5. Immediate safety concerns (if any)
            6. Recommended maintenance actions
            7. Priority level for addressing identified issues
            
            Provide specific, actionable recommendations based on the computer vision analysis results.
            Consider the machine context and typical failure modes for this type of equipment.
            """
            
            response = await self.llm_provider.get_best_response(prompt)
            return response
            
        except Exception as e:
            logger.error(f"LLaMA vision analysis error: {e}")
            return f"Visual analysis error: {e}"
# Enhanced Knowledge Base
class EnhancedMaintenanceKnowledgeBase:
    def __init__(self, db_path: str = "enhanced_maintenance.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Initialize enhanced SQLite database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Equipment table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS equipment (
                id TEXT PRIMARY KEY,
                name TEXT,
                type TEXT,
                model TEXT,
                manufacturer TEXT,
                installation_date DATE,
                last_maintenance DATE
            )
        ''')
        
        # Machine documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS machine_documents (
                doc_id TEXT PRIMARY KEY,
                machine_id TEXT,
                file_path TEXT,
                file_type TEXT,
                content TEXT,
                summary TEXT,
                upload_date TIMESTAMP,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Visual inspections table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS visual_inspections (
                inspection_id TEXT PRIMARY KEY,
                machine_id TEXT,
                file_path TEXT,
                file_type TEXT,
                defects_data TEXT,
                ai_analysis TEXT,
                timestamp TIMESTAMP,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Diagnosis sessions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS diagnosis_sessions (
                session_id TEXT PRIMARY KEY,
                machine_id TEXT,
                start_time TIMESTAMP,
                current_diagnosis TEXT,
                conversation_history TEXT,
                uploaded_files TEXT,
                status TEXT,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Enhanced maintenance history
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS maintenance_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                equipment_id TEXT,
                maintenance_date DATE,
                type TEXT,
                description TEXT,
                technician TEXT,
                cost REAL,
                downtime_hours INTEGER,
                supporting_documents TEXT,
                FOREIGN KEY (equipment_id) REFERENCES equipment (id)
            )
        ''')
        
        # Fault patterns table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS fault_patterns (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                equipment_type TEXT,
                symptoms TEXT,
                root_cause TEXT,
                solution TEXT,
                confidence_score REAL,
                supporting_evidence TEXT
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def add_machine_document(self, document: MachineDocument) -> bool:
        """Add machine document to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO machine_documents 
                (doc_id, machine_id, file_path, file_type, content, summary, upload_date)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                document.doc_id,
                document.machine_id,
                document.file_path,
                document.file_type,
                document.content,
                document.summary,
                document.upload_date
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding machine document: {e}")
            return False
    
    def get_machine_documents(self, machine_id: str) -> List[MachineDocument]:
        """Get all documents for a specific machine"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM machine_documents WHERE machine_id = ?
        ''', (machine_id,))
        
        results = cursor.fetchall()
        conn.close()
        
        documents = []
        for row in results:
            doc = MachineDocument(
                doc_id=row[0],
                machine_id=row[1],
                file_path=row[2],
                file_type=row[3],
                content=row[4],
                upload_date=datetime.fromisoformat(row[6]),
                summary=row[5] or ""
            )
            documents.append(doc)
        
        return documents
    
    def add_visual_inspection(self, inspection: VisualInspection) -> bool:
        """Add visual inspection to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO visual_inspections 
                (inspection_id, machine_id, file_path, file_type, defects_data, ai_analysis, timestamp)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                inspection.inspection_id,
                inspection.machine_id,
                inspection.file_path,
                inspection.file_type,
                json.dumps(inspection.defects_detected),
                inspection.ai_analysis,
                inspection.timestamp
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding visual inspection: {e}")
            return False
    
    def get_visual_inspections(self, machine_id: str) -> List[VisualInspection]:
        """Get all visual inspections for a machine"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM visual_inspections WHERE machine_id = ?
        ''', (machine_id,))
        
        results = cursor.fetchall()
        conn.close()
        
        inspections = []
        for row in results:
            inspection = VisualInspection(
                inspection_id=row[0],
                machine_id=row[1],
                file_path=row[2],
                file_type=row[3],
                defects_detected=json.loads(row[4]) if row[4] else [],
                timestamp=datetime.fromisoformat(row[6]),
                ai_analysis=row[5] or ""
            )
            inspections.append(inspection)
        
        return inspections
    
    def save_diagnosis_session(self, session: DiagnosisSession) -> bool:
        """Save diagnosis session"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO diagnosis_sessions 
                (session_id, machine_id, start_time, current_diagnosis, conversation_history, uploaded_files, status)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                session.session_id,
                session.machine_id,
                session.start_time,
                json.dumps(session.current_diagnosis.__dict__ if session.current_diagnosis else {}),
                json.dumps(session.conversation_history),
                json.dumps(session.uploaded_files),
                session.status
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error saving diagnosis session: {e}")
            return False
    
    def get_diagnosis_session(self, session_id: str) -> Optional[DiagnosisSession]:
        """Retrieve diagnosis session"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('SELECT * FROM diagnosis_sessions WHERE session_id = ?', (session_id,))
            result = cursor.fetchone()
            conn.close()
            
            if result:
                diagnosis_data = json.loads(result[3]) if result[3] else {}
                diagnosis = FaultDiagnosis(**diagnosis_data) if diagnosis_data else None
                
                return DiagnosisSession(
                    session_id=result[0],
                    machine_id=result[1],
                    start_time=datetime.fromisoformat(result[2]),
                    current_diagnosis=diagnosis,
                    conversation_history=json.loads(result[4]),
                    uploaded_files=json.loads(result[5]),
                    status=result[6]
                )
            return None
        except Exception as e:
            logger.error(f"Error retrieving diagnosis session: {e}")
            return None
    
    def add_equipment(self, equipment_data: Dict) -> bool:
        """Add new equipment to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO equipment (id, name, type, model, manufacturer, installation_date, last_maintenance)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                equipment_data['id'],
                equipment_data['name'],
                equipment_data['type'],
                equipment_data.get('model', 'Unknown'),
                equipment_data.get('manufacturer', 'Unknown'),
                equipment_data.get('installation_date', datetime.now().strftime('%Y-%m-%d')),
                equipment_data.get('last_maintenance', 'Never')
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding equipment: {e}")
            return False
    
    def get_all_equipment(self) -> List[Dict]:
        """Get all equipment from database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM equipment')
        equipment = cursor.fetchall()
        columns = [description[0] for description in cursor.description]
        conn.close()
        
        return [dict(zip(columns, row)) for row in equipment]

    def add_maintenance_record(self, maintenance_data: Dict) -> bool:
        """Add maintenance record"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT INTO maintenance_history 
                (equipment_id, maintenance_date, type, description, technician, cost, downtime_hours, supporting_documents)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                maintenance_data['equipment_id'],
                maintenance_data['maintenance_date'],
                maintenance_data['type'],
                maintenance_data['description'],
                maintenance_data.get('technician', 'Unknown'),
                maintenance_data.get('cost', 0.0),
                maintenance_data.get('downtime_hours', 0),
                json.dumps(maintenance_data.get('supporting_documents', []))
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding maintenance record: {e}")
            return False

# Enhanced Diagnostic Engine
class EnhancedDiagnosticEngine:
    def __init__(self, knowledge_base: EnhancedMaintenanceKnowledgeBase, llm_provider: EnhancedLLMProvider):
        self.knowledge_base = knowledge_base
        self.llm_provider = llm_provider
    
    async def comprehensive_diagnosis(self, machine_id: str, sensor_data: Optional[Dict] = None, visual_data: Optional[List[VisualInspection]] = None, user_description: str = "") -> FaultDiagnosis:
        """Perform comprehensive diagnosis using all available data"""
        
        # Get machine documents
        machine_docs = self.knowledge_base.get_machine_documents(machine_id)
        
        # Get visual inspections if not provided
        if visual_data is None:
            visual_data = self.knowledge_base.get_visual_inspections(machine_id)
        
        # Prepare comprehensive analysis prompt
        prompt = f"""
        COMPREHENSIVE MACHINE DIAGNOSIS REQUEST
        
        Machine ID: {machine_id}
        User Description: {user_description}
        
        AVAILABLE DATA:
        1. Machine Documentation: {len(machine_docs)} documents available
        2. Visual Inspections: {len(visual_data)} inspections available
        3. Sensor Data: {"Available" if sensor_data else "Not available"}
        
        ANALYSIS REQUEST:
        Based on the machine-specific documentation and visual inspection data, please provide:
        
        1. PRIMARY DIAGNOSIS:
           - Most likely root cause
           - Confidence level (0.0-1.0)
           - Severity assessment (low/medium/high/critical)
        
        2. SUPPORTING EVIDENCE:
           - Reference specific documentation sections
           - Cite visual inspection findings
           - Correlate with sensor anomalies
        
        3. RECOMMENDED DIAGNOSTIC STEPS:
           - Immediate actions to confirm diagnosis
           - Additional tests or inspections needed
           - Safety precautions
        
        4. ALTERNATIVE DIAGNOSES:
           - Other possible causes
           - How to differentiate between them
        
        5. REPAIR STRATEGY:
           - Step-by-step repair process
           - Required parts and tools
           - Estimated time and complexity
        
        Please provide specific, actionable recommendations based on the machine's actual documentation and inspection data.
        """
        
        # Get AI analysis with context
        ai_response = await self.llm_provider.analyze_with_context(prompt, machine_docs, visual_data)
        
        # Create unique fault ID using datetime formatting separately
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        fault_id = f"DIAG_{machine_id}_{timestamp_str}"
        
        # Parse response and create diagnosis
        diagnosis = FaultDiagnosis(
            fault_id=fault_id,
            machine_id=machine_id,
            severity=self._extract_severity(ai_response),
            confidence=self._extract_confidence(ai_response),
            description=user_description,
            root_cause=self._extract_root_cause(ai_response),
            recommended_actions=self._extract_recommended_actions(ai_response),
            estimated_downtime=self._estimate_downtime(ai_response),
            supporting_evidence=self._extract_supporting_evidence(ai_response, machine_docs, visual_data)
        )
        
        return diagnosis
    
    def _extract_severity(self, ai_response: str) -> str:
        """Extract severity from AI response"""
        response_lower = ai_response.lower()
        if 'critical' in response_lower:
            return 'critical'
        elif 'high' in response_lower:
            return 'high'
        elif 'medium' in response_lower:
            return 'medium'
        else:
            return 'low'

    def _extract_confidence(self, ai_response: str) -> float:
        """Extract confidence level from AI response"""
        # Look for confidence patterns
        confidence_patterns = [
            r'confidence[:\s]+(\d+\.?\d*)%',
            r'(\d+\.?\d*)%\s+confidence',
            r'(\d+\.?\d*)\s+confidence'
        ]
        
        for pattern in confidence_patterns:
            match = re.search(pattern, ai_response.lower())
            if match:
                value = float(match.group(1))
                if value > 1.0:  # If percentage
                    return min(value / 100.0, 1.0)
                else:
                    return min(value, 1.0)
        
        return 0.7  # Default confidence

    def _extract_root_cause(self, ai_response: str) -> str:
        """Extract root cause from AI response"""
        lines = ai_response.split('\n')
        for line in lines:
            if 'root cause' in line.lower() or 'primary diagnosis' in line.lower():
                return line.strip()
        
        # Fallback: return first few sentences
        sentences = ai_response.split('.')[:3]
        return '. '.join(sentences) + '.'

    def _extract_recommended_actions(self, ai_response: str) -> List[str]:
        """Extract recommended actions from AI response"""
        actions = []
        lines = ai_response.split('\n')
        in_actions_section = False
        
        for line in lines:
            line = line.strip()
            if 'recommended' in line.lower() and ('action' in line.lower() or 'step' in line.lower()):
                in_actions_section = True
                continue
            
            if in_actions_section:
                if line.startswith(('-', '*', '•')) or re.match(r'^\d+\.', line):
                    actions.append(line.lstrip('- *•0123456789. '))
                elif line and not line.lower().startswith(('alternative', 'repair strategy', 'supporting')):
                    continue
                else:
                    break
        
        return actions[:10]  # Limit to 10 actions

    def _estimate_downtime(self, ai_response: str) -> int:
        """Estimate downtime from AI response"""
        # Look for time patterns
        time_patterns = [
            r'(\d+)\s*hours?',
            r'(\d+)\s*days?',
            r'(\d+)\s*minutes?'
        ]
        
        for pattern in time_patterns:
            matches = re.findall(pattern, ai_response.lower())
            if matches:
                # Convert to hours
                value = int(matches[0])
                if 'day' in pattern:
                    return value * 24
                elif 'minute' in pattern:
                    return max(1, value // 60)
                else:
                    return value
        
        # Default based on severity
        severity = self._extract_severity(ai_response)
        if severity == 'critical':
            return 24
        elif severity == 'high':
            return 8
        elif severity == 'medium':
            return 4
        else:
            return 1

    def _extract_supporting_evidence(self, ai_response: str, machine_docs: List[MachineDocument], visual_data: List[VisualInspection]) -> List[str]:
        """Extract supporting evidence references"""
        evidence = []
        
        # Add document references
        for doc in machine_docs:
            if doc.file_path.lower() in ai_response.lower():
                evidence.append(f"Document: {doc.file_path}")
        
        # Add visual inspection references
        for inspection in visual_data:
            if len(inspection.defects_detected) > 0:
                evidence.append(f"Visual inspection: {len(inspection.defects_detected)} defects detected in {inspection.file_path}")
        
        # Add any other evidence mentioned in response
        lines = ai_response.split('\n')
        for line in lines:
            if 'evidence' in line.lower() and line.strip():
                evidence.append(line.strip())
        
        return evidence

    async def start_interactive_diagnosis(self, machine_id: str) -> DiagnosisSession:
        """Start interactive diagnosis session"""
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        session_id = f"SESSION_{machine_id}_{timestamp_str}"
        
        session = DiagnosisSession(
            session_id=session_id,
            machine_id=machine_id,
            start_time=datetime.now(),
            current_diagnosis=None,
            conversation_history=[],
            uploaded_files=[],
            status="active"
        )
        
        # Save session
        self.knowledge_base.save_diagnosis_session(session)
        
        return session

    async def continue_diagnosis(self, session_id: str, user_input: str, file_paths: List[str] = None) -> Tuple[str, DiagnosisSession]:
        """Continue interactive diagnosis"""
        session = self.knowledge_base.get_diagnosis_session(session_id)
        if not session:
            return "Session not found", None
        
        # Add user input to conversation
        session.conversation_history.append({
            "role": "user",
            "content": user_input,
            "timestamp": datetime.now().isoformat()
        })
        
        # Process any new files
        if file_paths:
            for file_path in file_paths:
                if file_path not in session.uploaded_files:
                    session.uploaded_files.append(file_path)
        
        # Get machine context
        machine_docs = self.knowledge_base.get_machine_documents(session.machine_id)
        visual_data = self.knowledge_base.get_visual_inspections(session.machine_id)
        
        # Build conversation context
        conversation_context = "\n".join([
            f"{msg['role']}: {msg['content']}" for msg in session.conversation_history[-5:]  # Last 5 messages
        ])
        
        prompt = f"""
        INTERACTIVE MAINTENANCE DIAGNOSIS - Session {session_id}
        Machine ID: {session.machine_id}
        
        CONVERSATION HISTORY:
        {conversation_context}
        
        LATEST USER INPUT: {user_input}
        
        As an expert maintenance engineer, continue this diagnostic conversation.
        Consider the machine documentation and previous conversation context.
        Ask follow-up questions if needed, or provide specific guidance based on the information available.
        
        If you have enough information, provide a preliminary diagnosis with confidence level.
        """
        
        ai_response = await self.llm_provider.analyze_with_context(prompt, machine_docs, visual_data)
        
        # Add AI response to conversation
        session.conversation_history.append({
            "role": "assistant",
            "content": ai_response,
            "timestamp": datetime.now().isoformat()
        })
        
        # Update session
        self.knowledge_base.save_diagnosis_session(session)
        
        return ai_response, session

# Enhanced Smart Maintenance System
class EnhancedSmartMaintenanceSystem:
    def __init__(self):
        self.knowledge_base = EnhancedMaintenanceKnowledgeBase()
        self.llm_provider = EnhancedLLMProvider()
        self.diagnostic_engine = EnhancedDiagnosticEngine(self.knowledge_base, self.llm_provider)
        self.document_processor = DocumentProcessor()
        self.vision_processor = EnhancedVisionProcessor(self.llm_provider)
    
    async def test_connection(self):
        """Test the LLM connection"""
        test_response = await self.llm_provider.call_llama("Test connection: What are common causes of bearing failure?")
        return test_response
    
    async def close(self):
        """Cleanup system resources"""
        await self.llm_provider.close()

# Example usage and testing
async def main():
    """Example usage of the Enhanced Smart Maintenance System with LLaMA"""
    system = EnhancedSmartMaintenanceSystem()
    
    try:
        # Test LLaMA connection
        print("Testing LLaMA API connection...")
        test_response = await system.test_connection()
        print(f"LLaMA Test Response: {test_response[:200]}...")
        
        # Add some sample equipment
        sample_equipment = {
            'id': 'PUMP_001',
            'name': 'Main Water Pump',
            'type': 'Centrifugal Pump',
            'model': 'CP-500',
            'manufacturer': 'Industrial Pumps Inc.',
            'installation_date': '2020-01-15',
            'last_maintenance': '2024-01-15'
        }
        
        system.knowledge_base.add_equipment(sample_equipment)
        print("Sample equipment added")
        
        # Example: Start diagnosis
        diagnosis = await system.diagnostic_engine.comprehensive_diagnosis(
            machine_id='PUMP_001',
            user_description="Pump is making unusual noise and vibrating more than normal"
        )
        print(f"Diagnosis completed: {diagnosis.fault_id}")
        print(f"Severity: {diagnosis.severity}")
        print(f"Root cause: {diagnosis.root_cause}")
        
    except Exception as e:
        print(f"Error in main: {e}")
    
    finally:
        await system.close()

if __name__ == "__main__":
    asyncio.run(main())
