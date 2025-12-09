import asyncio
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass
from enum import Enum
import aiohttp
from dotenv import load_dotenv
import PyPDF2
import docx
from pathlib import Path
import hashlib
import shutil
load_dotenv()

class LLMConfig:
    """Configuration for free LLM providers"""
    # Free providers only
    GROQ_API_KEY = os.getenv("GROQ_API_KEY", "your-groq-api-key")
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "your-gemini-api-key")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enhanced Data Models
@dataclass
class SensorData:
    sensor_id: str
    timestamp: datetime
    temperature: float
    vibration: float
    pressure: float
    humidity: float
    sound_level: float

@dataclass
class MachineDocument:
    doc_id: str
    machine_id: str
    file_path: str
    file_type: str
    content: str
    upload_date: datetime
    summary: str = ""

@dataclass
class VisualInspection:
    inspection_id: str
    machine_id: str
    file_path: str
    file_type: str  # image or video
    defects_detected: List[Dict]
    timestamp: datetime
    ai_analysis: str = ""

@dataclass
class FaultDiagnosis:
    fault_id: str
    machine_id: str
    severity: str
    confidence: float
    description: str
    root_cause: str
    recommended_actions: List[str]
    estimated_downtime: int
    supporting_evidence: List[str]  # References to documents/images
    conversation_history: Optional[List[Dict]] = None

@dataclass
class DiagnosisSession:
    session_id: str
    machine_id: str
    start_time: datetime
    current_diagnosis: Optional[FaultDiagnosis]
    conversation_history: List[Dict]
    uploaded_files: List[str]
    status: str  # active, resolved, escalated

class MaintenanceType(Enum):
    PREVENTIVE = "preventive"
    CORRECTIVE = "corrective"
    PREDICTIVE = "predictive"
    EMERGENCY = "emergency"

@dataclass
class MaintenanceTask:
    task_id: str
    equipment_id: str
    task_type: MaintenanceType
    priority: int
    description: str
    scheduled_date: datetime
    estimated_duration: int
    required_parts: List[str]
    assigned_technician: Optional[str] = None

# Enhanced Document Processing
class DocumentProcessor:
    def __init__(self, storage_dir: str = "uploads"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
    def save_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded file and return new path"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def process_document(self, file_path: str, machine_id: str) -> MachineDocument:
        """Process uploaded document and extract content"""
        file_path_obj = Path(file_path)
        file_type = file_path_obj.suffix.lower()
        
        # Extract text based on file type
        if file_type == '.pdf':
            content = self.extract_text_from_pdf(file_path)
        elif file_type == '.docx':
            content = self.extract_text_from_docx(file_path)
        elif file_type in ['.txt', '.md']:
            content = self.extract_text_from_txt(file_path)
        else:
            content = f"Unsupported file type: {file_type}"
        
        # Generate document ID
        doc_id = hashlib.md5(f"{machine_id}_{file_path_obj.name}_{datetime.now()}".encode()).hexdigest()[:12]
        
        return MachineDocument(
            doc_id=doc_id,
            machine_id=machine_id,
            file_path=file_path,
            file_type=file_type,
            content=content,
            upload_date=datetime.now()
        )

class EnhancedLLMProvider:
    def __init__(self):
        self.current_provider = "auto"  # auto, groq, gemini, huggingface, local
        self.init_providers()
    
    def set_provider(self, provider: str):
        """Set specific LLM provider"""
        valid_providers = ["auto", "groq", "gemini", "huggingface", "local"]
        if provider.lower() in valid_providers:
            self.current_provider = provider.lower()
            print(f"LLM provider set to: {self.current_provider}")
        else:
            print(f"Invalid provider. Choose from: {valid_providers}")
    
    def get_current_provider(self):
        """Get current LLM provider"""
        return self.current_provider
    
    def init_providers(self):
        """Initialize free LLM providers only"""
        # Groq (Free tier available)
        self.groq_configured = (LLMConfig.GROQ_API_KEY and 
                              LLMConfig.GROQ_API_KEY != "your-groq-api-key")
        
        # Gemini (Free tier available)
        self.gemini_configured = (LLMConfig.GEMINI_API_KEY and 
                                LLMConfig.GEMINI_API_KEY != "your-gemini-api-key")
        
        logger.info(f"Free providers configured: "
                   f"Groq={self.groq_configured}, "
                   f"Gemini={self.gemini_configured}")
    
    async def call_huggingface(self, prompt: str) -> str:
        """Call Hugging Face API (free tier available)"""
        try:
            if not self.hf_configured:
                return "Hugging Face API not configured"
            
            headers = {
                "Authorization": f"Bearer {LLMConfig.HUGGINGFACE_API_KEY}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": 512,
                    "temperature": 0.7,
                    "return_full_text": False
                }
            }
            
            # Try multiple free models in order of preference
            models = [
                "microsoft/DialoGPT-medium",  # Smaller, more reliable model
                "google/flan-t5-base",        # Alternative free model
                "facebook/blenderbot-400M-distill"  # Another option
            ]
            
            for model in models:
                model_url = f"https://api-inference.huggingface.co/models/{model}"
                
                try:
                    async with aiohttp.ClientSession() as session:
                        async with session.post(
                            model_url,
                            headers=headers,
                            json=payload,
                            timeout=aiohttp.ClientTimeout(total=30)
                        ) as response:
                            if response.status == 200:
                                data = await response.json()
                                
                                # Handle different response formats
                                if isinstance(data, list) and len(data) > 0:
                                    if 'generated_text' in data[0]:
                                        return data[0]['generated_text']
                                    elif 'translation_text' in data[0]:
                                        return data[0]['translation_text']
                                    else:
                                        return str(data[0])
                                elif isinstance(data, dict):
                                    if 'generated_text' in data:
                                        return data['generated_text']
                                    else:
                                        return str(data)
                                return str(data)
                            elif response.status == 503:
                                # Model loading, try next model
                                continue
                            else:
                                # Try next model on error
                                continue
                                
                except Exception as e:
                    logger.warning(f"HF model {model} failed: {e}")
                    continue
            
            # If all models fail, return fallback
            return "Hugging Face API: All models unavailable. Using local fallback."
                        
        except Exception as e:
            logger.error(f"HF API error: {e}")
            return f"Error calling Hugging Face: {e}"
    
    async def call_groq(self, prompt: str) -> str:
        """Call Groq API (free tier available)"""
        try:
            if not self.groq_configured:
                return "Groq API not configured"
            
            headers = {
                "Authorization": f"Bearer {LLMConfig.GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            
            # Try different free models
            models = [
                "llama3-8b-8192",    # Primary model
                "mixtral-8x7b-32768", # Alternative model
                "gemma-7b-it"        # Another option
            ]
            
            for model in models:
                try:
                    payload = {
                        "model": model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": 1024,
                        "temperature": 0.7
                    }
                    
                    async with aiohttp.ClientSession() as session:
                        async with session.post(
                            "https://api.groq.com/openai/v1/chat/completions",
                            headers=headers,
                            json=payload,
                            timeout=aiohttp.ClientTimeout(total=30)
                        ) as response:
                            if response.status == 200:
                                data = await response.json()
                                return data["choices"][0]["message"]["content"].strip()
                            elif response.status == 429:
                                # Rate limit, try next model
                                continue
                            else:
                                # Try next model on error
                                continue
                                
                except Exception as e:
                    logger.warning(f"Groq model {model} failed: {e}")
                    continue
            
            return "Groq API: All models unavailable"
                        
        except Exception as e:
            logger.error(f"Groq API error: {e}")
            return f"Error calling Groq: {e}"
    
    async def call_gemini(self, prompt: str) -> str:
        """Call Gemini API (free tier available)"""
        try:
            if not self.gemini_configured:
                return "Gemini API not configured"
            
            # Try different Gemini models
            models = [
                "gemini-pro",           # Primary model
                "gemini-1.5-flash"      # Alternative model
            ]
            
            for model in models:
                try:
                    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={LLMConfig.GEMINI_API_KEY}"
                    
                    payload = {
                        "contents": [{
                            "parts": [{"text": prompt}]
                        }],
                        "generationConfig": {
                            "temperature": 0.7,
                            "maxOutputTokens": 1024
                        }
                    }
                    
                    async with aiohttp.ClientSession() as session:
                        async with session.post(
                            url,
                            json=payload,
                            timeout=aiohttp.ClientTimeout(total=30)
                        ) as response:
                            if response.status == 200:
                                data = await response.json()
                                if 'candidates' in data and len(data['candidates']) > 0:
                                    content = data['candidates'][0]['content']['parts'][0]['text']
                                    return content.strip()
                                return "No response generated"
                            elif response.status == 429:
                                # Rate limit, try next model
                                continue
                            else:
                                # Try next model on error
                                continue
                                
                except Exception as e:
                    logger.warning(f"Gemini model {model} failed: {e}")
                    continue
            
            return "Gemini API: All models unavailable"
                        
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            return f"Error calling Gemini: {e}"
    
    async def call_local_fallback(self, prompt: str) -> str:
        """Enhanced local fallback with better maintenance knowledge"""
        maintenance_keywords = {
            'bearing': {
                'response': "BEARING ANALYSIS: Common bearing failure causes include inadequate lubrication (40%), misalignment (25%), contamination (15%), overloading (10%), and fatigue (10%). Recommended actions: 1) Check lubrication levels and quality, 2) Inspect shaft alignment, 3) Look for contamination sources, 4) Verify load conditions, 5) Replace if excessive wear detected.",
                'symptoms': ['noise', 'vibration', 'heat', 'play']
            },
            'pump': {
                'response': "PUMP ANALYSIS: Common pump issues include cavitation (suction problems), impeller wear, mechanical seal failure, and misalignment. Check: 1) Suction pressure and NPSH, 2) Impeller condition and clearances, 3) Mechanical seal integrity, 4) Shaft alignment, 5) Motor coupling condition.",
                'symptoms': ['noise', 'vibration', 'leak', 'flow', 'pressure']
            },
            'motor': {
                'response': "MOTOR ANALYSIS: Common motor problems include overheating, bearing failure, insulation breakdown, and rotor issues. Check: 1) Operating temperature, 2) Vibration levels, 3) Electrical connections and insulation, 4) Current draw and power factor, 5) Bearing condition.",
                'symptoms': ['heat', 'noise', 'vibration', 'current', 'speed']
            },
            'vibration': {
                'response': "VIBRATION ANALYSIS: Excessive vibration typically indicates imbalance, misalignment, bearing wear, or loose components. Systematic approach: 1) Measure vibration amplitude and frequency, 2) Check for loose bolts/foundations, 3) Verify alignment, 4) Inspect bearings, 5) Check for resonance conditions.",
                'causes': ['imbalance', 'misalignment', 'bearing', 'looseness']
            },
            'noise': {
                'response': "NOISE ANALYSIS: Unusual equipment noise sources include bearing wear, cavitation, loose parts, and misalignment. Investigation steps: 1) Identify noise frequency and type, 2) Locate noise source, 3) Check for cavitation conditions, 4) Inspect bearings and connections, 5) Verify proper installation.",
                'types': ['grinding', 'squealing', 'knocking', 'humming']
            },
            'leak': {
                'response': "LEAK ANALYSIS: Common leak sources include seal failure, gasket degradation, pipe joints, and valve packing. Systematic inspection: 1) Identify leak location and type, 2) Check seal condition and installation, 3) Verify proper torque on connections, 4) Inspect gasket materials, 5) Consider pressure and temperature effects.",
                'locations': ['seals', 'gaskets', 'joints', 'valves']
            },
            'temperature': {
                'response': "TEMPERATURE ANALYSIS: Abnormal temperatures indicate friction, insufficient cooling, overloading, or electrical issues. Check: 1) Cooling system operation, 2) Lubrication adequacy, 3) Load conditions, 4) Electrical connections, 5) Ambient conditions.",
                'issues': ['overheating', 'overcooling', 'fluctuation']
            },
            'pressure': {
                'response': "PRESSURE ANALYSIS: Pressure issues in systems indicate blockages, leaks, pump problems, or control valve issues. Investigation: 1) Check pressure at multiple points, 2) Verify pump operation, 3) Inspect filters and strainers, 4) Check control valve operation, 5) Look for system leaks.",
                'types': ['low', 'high', 'fluctuating', 'dropping']
            }
        }
        
        prompt_lower = prompt.lower()
        
        # Find the most relevant keyword
        best_match = None
        max_score = 0
        
        for keyword, data in maintenance_keywords.items():
            score = 0
            if keyword in prompt_lower:
                score += 10
            
            # Check for related symptoms/causes
            if 'symptoms' in data:
                for symptom in data['symptoms']:
                    if symptom in prompt_lower:
                        score += 5
            if 'causes' in data:
                for cause in data['causes']:
                    if cause in prompt_lower:
                        score += 5
            if 'types' in data:
                for ptype in data['types']:
                    if ptype in prompt_lower:
                        score += 5
            if 'locations' in data:
                for location in data['locations']:
                    if location in prompt_lower:
                        score += 5
            if 'issues' in data:
                for issue in data['issues']:
                    if issue in prompt_lower:
                        score += 5
            
            if score > max_score:
                max_score = score
                best_match = data['response']
        
        if best_match:
            return f"MAINTENANCE ANALYSIS (Local Knowledge Base): {best_match}"
        
        # General maintenance advice if no specific match
        if any(word in prompt_lower for word in ['problem', 'issue', 'failure', 'fault', 'broken']):
            return """GENERAL MAINTENANCE ANALYSIS (Local Knowledge Base): 
            For effective troubleshooting: 1) Document all symptoms and operating conditions, 2) Check recent maintenance history, 3) Verify proper operating procedures, 4) Inspect for obvious physical damage, 5) Measure key parameters (temperature, pressure, vibration), 6) Consult equipment manuals, 7) Consider environmental factors. Always prioritize safety and follow lockout/tagout procedures."""
        
        return "MAINTENANCE ANALYSIS (Local Knowledge Base): Please provide more specific information about the equipment type, symptoms observed, and operating conditions for detailed analysis."
    
    async def get_best_response(self, prompt: str) -> str:
        """Get response from specified provider or try providers in order"""
        # Enhanced prompt for better maintenance responses
        enhanced_prompt = f"""
        You are an expert industrial maintenance engineer with 20+ years of experience in troubleshooting mechanical, electrical, and process equipment. 
        
        User Query: {prompt}
        
        Please provide a detailed, technical response that includes:
        1. Root cause analysis
        2. Systematic troubleshooting approach
        3. Safety considerations
        4. Specific recommendations
        5. Preventive measures
        
        Focus on practical, actionable advice based on industry best practices.
        """
        
        if self.current_provider == "groq" and self.groq_configured:
            response = await self.call_groq(enhanced_prompt)
            if not any(error in response.lower() for error in ['error', 'unavailable', 'not configured']):
                return response
        elif self.current_provider == "gemini" and self.gemini_configured:
            response = await self.call_gemini(enhanced_prompt)
            if not any(error in response.lower() for error in ['error', 'unavailable', 'not configured']):
                return response
        elif self.current_provider == "local":
            return await self.call_local_fallback(enhanced_prompt)
        elif self.current_provider == "auto":
            # Try providers in order of preference
            providers = [
                ("Groq", self.call_groq) if self.groq_configured else None,
                ("Gemini", self.call_gemini) if self.gemini_configured else None,
                ("Local Fallback", self.call_local_fallback)
            ]
            
            # Filter out None providers
            providers = [p for p in providers if p is not None]
            
            for provider_name, provider_func in providers:
                try:
                    logger.info(f"Trying {provider_name} provider...")
                    response = await provider_func(enhanced_prompt)
                    
                    # Check if response indicates an error
                    if not any(error_term in response.lower() for error_term in 
                              ['error', 'quota', 'exceeded', 'not configured', 'unavailable']):
                        logger.info(f"Success with {provider_name}")
                        return response
                    else:
                        logger.warning(f"{provider_name} failed: {response[:100]}...")
                        
                except Exception as e:
                    logger.error(f"{provider_name} provider failed: {e}")
                    continue
            
            # If all providers fail, return local fallback
            return await self.call_local_fallback(enhanced_prompt)
        else:
            return "Invalid provider configuration"

    async def test_connection_detailed(self):
        """Test the free LLM provider connections with detailed status"""
        test_prompt = "What are the top 3 causes of pump cavitation? Respond briefly."
        
        results = {}
        
        # Test each provider individually
        if self.groq_configured:
            try:
                response = await self.call_groq(test_prompt)
                results['groq'] = 'OK' if not any(err in response.lower() for err in ['error', 'unavailable']) else 'FAILED'
            except:
                results['groq'] = 'FAILED'
        else:
            results['groq'] = 'NOT_CONFIGURED'
        
        if self.gemini_configured:
            try:
                response = await self.call_gemini(test_prompt)
                results['gemini'] = 'OK' if not any(err in response.lower() for err in ['error', 'unavailable']) else 'FAILED'
            except:
                results['gemini'] = 'FAILED'
        else:
            results['gemini'] = 'NOT_CONFIGURED'
        
        # Local fallback is always available
        results['local'] = 'OK'
        
        return results

class EnhancedSmartMaintenanceSystem:
    """Enhanced Smart Maintenance Diagnostic System"""
    
    def __init__(self):
        self.llm_provider = EnhancedLLMProvider()
        self.document_processor = DocumentProcessor()
        self.equipment_database = {}
        self.diagnosis_history = []
        self.current_session = None
    
    def display_menu(self):
        """Display main menu"""
        print("\n" + "="*50)
        print("Smart Maintenance Diagnostic System")
        print("="*50)
        print("1. Add Equipment")
        print("2. Start Diagnosis Session")
        print("3. Upload Documents")
        print("4. Upload Visual Inspection")
        print("5. Continue Active Diagnosis")
        print("6. View Equipment List")
        print("7. View Diagnosis History")
        print("8. Test LLM Connection")
        print("9. Change LLM Provider")
        print("10. Chat with AI")
        print("11. Exit")
        print("="*50)
    
    async def add_equipment_interactive(self):
        """Add equipment interactively"""
        print("\n--- Add Equipment ---")
        machine_id = input("Enter machine ID: ").strip()
        machine_name = input("Enter machine name: ").strip()
        machine_type = input("Enter machine type: ").strip()
        
        self.equipment_database[machine_id] = {
            'name': machine_name,
            'type': machine_type,
            'added_date': datetime.now()
        }
        print(f"Equipment {machine_id} added successfully!")
    
    async def start_diagnosis_interactive(self):
        """Start diagnosis session interactively"""
        print("\n--- Start Diagnosis Session ---")
        
        if not self.equipment_database:
            print("No equipment registered. Please add equipment first.")
            return
        
        print("Available equipment:")
        for machine_id, info in self.equipment_database.items():
            print(f"- {machine_id}: {info['name']} ({info['type']})")
        
        machine_id = input("Enter machine ID for diagnosis: ").strip()
        
        if machine_id not in self.equipment_database:
            print("Equipment not found!")
            return
        
        session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        self.current_session = DiagnosisSession(
            session_id=session_id,
            machine_id=machine_id,
            start_time=datetime.now(),
            current_diagnosis=None,
            conversation_history=[],
            uploaded_files=[],
            status="active"
        )
        
        print(f"Diagnosis session {session_id} started for {machine_id}")
        await self.continue_interactive_diagnosis(session_id)
    
    async def upload_documents_interactive(self):
        """Upload documents interactively"""
        print("\n--- Upload Documents ---")
        
        if not self.equipment_database:
            print("No equipment registered. Please add equipment first.")
            return
        
        print("Available equipment:")
        for machine_id, info in self.equipment_database.items():
            print(f"- {machine_id}: {info['name']}")
        
        machine_id = input("Enter machine ID: ").strip()
        file_path = input("Enter file path: ").strip()
        
        if machine_id not in self.equipment_database:
            print("Equipment not found!")
            return
        
        if not os.path.exists(file_path):
            print("File not found!")
            return
        
        try:
            doc = self.document_processor.process_document(file_path, machine_id)
            print(f"Document processed successfully! Doc ID: {doc.doc_id}")
            print(f"Content length: {len(doc.content)} characters")
        except Exception as e:
            print(f"Error processing document: {e}")
    
    async def upload_visual_inspection_interactive(self):
        """Upload visual inspection interactively"""
        print("\n--- Upload Visual Inspection ---")
        print("Visual inspection upload simulated (image processing not implemented)")
    
    async def continue_interactive_diagnosis(self, session_id: str):
        """Continue interactive diagnosis"""
        print(f"\n--- Diagnosis Session: {session_id} ---")
        
        while True:
            user_input = input("\nDescribe the issue or ask a question (type 'exit' to end session): ").strip()
            
            if user_input.lower() == 'exit':
                if self.current_session:
                    self.current_session.status = "resolved"
                    self.diagnosis_history.append(self.current_session)
                    self.current_session = None
                print("Diagnosis session ended.")
                break
            
            if not user_input:
                continue
            
            # Get AI response
            try:
                response = await self.llm_provider.get_best_response(user_input)
                print(f"\nAI Diagnosis: {response}")
                
                # Add to conversation history
                if self.current_session:
                    self.current_session.conversation_history.append({
                        'user': user_input,
                        'ai': response,
                        'timestamp': datetime.now()
                    })
                
            except Exception as e:
                print(f"Error getting AI response: {e}")
    
    def view_equipment_list(self):
        """View equipment list"""
        print("\n--- Equipment List ---")
        
        if not self.equipment_database:
            print("No equipment registered.")
            return
        
        for machine_id, info in self.equipment_database.items():
            print(f"ID: {machine_id}")
            print(f"  Name: {info['name']}")
            print(f"  Type: {info['type']}")
            print(f"  Added: {info['added_date'].strftime('%Y-%m-%d %H:%M:%S')}")
            print()
    
    def view_diagnosis_history(self):
        """View diagnosis history"""
        print("\n--- Diagnosis History ---")
        
        if not self.diagnosis_history:
            print("No diagnosis history available.")
            return
        
        for session in self.diagnosis_history:
            print(f"Session: {session.session_id}")
            print(f"  Machine: {session.machine_id}")
            print(f"  Start Time: {session.start_time.strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"  Status: {session.status}")
            print(f"  Conversations: {len(session.conversation_history)}")
            print()
    
    async def test_connection(self):
        """Test LLM connection with detailed status"""
        print("Testing LLM providers...")
        results = await self.llm_provider.test_connection_detailed()
        
        print("\nLLM Provider Status:")
        print("-" * 30)
        for provider, status in results.items():
            status_symbol = "OK" if status == 'OK' else ("FAILED" if status == 'FAILED' else "NOT_CONFIGURED")
            print(f"{provider.capitalize()}: {status_symbol}")
        
        # Get a working response
        working_providers = [p for p, s in results.items() if s == 'OK']
        if working_providers:
            print(f"\nTesting response from {working_providers[0]}...")
            test_response = await self.llm_provider.get_best_response("What are common signs of bearing failure?")
            return test_response
        else:
            return "No working providers available"
    
    def change_llm_provider(self):
        """Change LLM provider"""
        print("\n--- Change LLM Provider ---")
        print("Available providers:")
        print("1. Auto (try all providers)")
        print("2. Groq")
        print("3. Gemini")
        print("4. Local Fallback")
        
        choice = input("Select provider (1-4): ").strip()
        
        provider_map = {
            '1': 'auto',
            '2': 'groq',
            '3': 'gemini',
            '4': 'local'
        }
        
        if choice in provider_map:
            self.llm_provider.set_provider(provider_map[choice])
        else:
            print("Invalid choice.")
    
    async def chat_with_ai(self):
        """Free chat with AI"""
        print("\n--- Chat with AI ---")
        print("Type 'exit' to return to main menu")
        
        while True:
            user_input = input("\nYou: ").strip()
            
            if user_input.lower() == 'exit':
                break
            
            if not user_input:
                continue
            
            try:
                response = await self.llm_provider.get_best_response(user_input)
                print(f"AI: {response}")
            except Exception as e:
                print(f"Error: {e}")

    async def run_interactive_system(self):
        """Main interactive system loop"""
        print("Starting Smart Maintenance Diagnostic System...")
        
        # Test connection on startup
        try:
            test_response = await self.test_connection()
            print("LLM Connection successful!")
        except Exception as e:
            print(f"LLM Connection warning: {e}")
        
        while True:
            try:
                self.display_menu()
                choice = input("Select option (1-11): ").strip()
                
                if choice == "1":
                    await self.add_equipment_interactive()
                elif choice == "2":
                    await self.start_diagnosis_interactive()
                elif choice == "3":
                    await self.upload_documents_interactive()
                elif choice == "4":
                    await self.upload_visual_inspection_interactive()
                elif choice == "5":
                    if self.current_session:
                        await self.continue_interactive_diagnosis(self.current_session.session_id)
                    else:
                        print("No active diagnosis session.")
                elif choice == "6":
                    self.view_equipment_list()
                elif choice == "7":
                    self.view_diagnosis_history()
                elif choice == "8":
                    print("Testing LLM connection...")
                    test_response = await self.test_connection()
                    print(f"Response: {test_response[:300]}...")
                elif choice == "9":
                    self.change_llm_provider()
                elif choice == "10":
                    await self.chat_with_ai()
                elif choice == "11":
                    print("Goodbye!")
                    break
                else:
                    print("Invalid option. Please select 1-11.")
                
                input("\nPress Enter to continue...")
                
            except KeyboardInterrupt:
                print("\n\nSystem interrupted. Goodbye!")
                break
            except Exception as e:
                print(f"Unexpected error: {e}")
                input("Press Enter to continue...")
    
    async def close(self):
        """Cleanup system resources"""
        # Close any open sessions or resources
        print("Cleaning up resources...")

# Interactive main function
async def main():
    """Interactive Smart Maintenance System"""
    system = EnhancedSmartMaintenanceSystem()
    
    try:
        await system.run_interactive_system()
    except Exception as e:
        print(f"System error: {e}")
    finally:
        await system.close()

if __name__ == "__main__":
    asyncio.run(main())