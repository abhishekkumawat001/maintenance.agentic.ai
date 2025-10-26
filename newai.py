import asyncio
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict
from enum import Enum
import aiohttp
from dotenv import load_dotenv
import PyPDF2
import docx
from pathlib import Path
import hashlib
import shutil
import json
from concurrent.futures import ThreadPoolExecutor
import threading

# Flask imports
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import uuid

load_dotenv()

class LLMConfig:
    """Configuration for free LLM providers"""
    # Free providers only
    GROQ_API_KEY = os.getenv("GROQ_API_KEY", "your-groq-api-key")
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "your-gemini-api-key")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enhanced Data Models
@dataclass
class SensorData:
    sensor_id: str
    timestamp: datetime
    temperature: float
    vibration: float
    pressure: float
    humidity: float
    sound_level: float

@dataclass
class MachineDocument:
    doc_id: str
    machine_id: str
    file_path: str
    file_type: str
    content: str
    upload_date: datetime
    summary: str = ""

@dataclass
class VisualInspection:
    inspection_id: str
    machine_id: str
    file_path: str
    file_type: str  # image or video
    defects_detected: List[Dict]
    timestamp: datetime
    ai_analysis: str = ""

@dataclass
class FaultDiagnosis:
    fault_id: str
    machine_id: str
    severity: str
    confidence: float
    description: str
    root_cause: str
    recommended_actions: List[str]
    estimated_downtime: int
    supporting_evidence: List[str]  # References to documents/images
    conversation_history: Optional[List[Dict]] = None

@dataclass
class DiagnosisSession:
    session_id: str
    machine_id: str
    start_time: datetime
    current_diagnosis: Optional[FaultDiagnosis]
    conversation_history: List[Dict]
    uploaded_files: List[str]
    status: str  # active, resolved, escalated

class MaintenanceType(Enum):
    PREVENTIVE = "preventive"
    CORRECTIVE = "corrective"
    PREDICTIVE = "predictive"
    EMERGENCY = "emergency"

@dataclass
class MaintenanceTask:
    task_id: str
    equipment_id: str
    task_type: MaintenanceType
    priority: int
    description: str
    scheduled_date: datetime
    estimated_duration: int
    required_parts: List[str]
    assigned_technician: Optional[str] = None

# Enhanced Document Processing
class DocumentProcessor:
    def __init__(self, storage_dir: str = "uploads"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
    def save_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded file and return new path"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def process_document(self, file_path: str, machine_id: str) -> MachineDocument:
        """Process uploaded document and extract content"""
        file_path_obj = Path(file_path)
        file_type = file_path_obj.suffix.lower()
        
        # Extract text based on file type
        if file_type == '.pdf':
            content = self.extract_text_from_pdf(file_path)
        elif file_type == '.docx':
            content = self.extract_text_from_docx(file_path)
        elif file_type in ['.txt', '.md']:
            content = self.extract_text_from_txt(file_path)
        else:
            content = f"Unsupported file type: {file_type}"
        
        # Generate document ID
        doc_id = hashlib.md5(f"{machine_id}_{file_path_obj.name}_{datetime.now()}".encode()).hexdigest()[:12]
        
        return MachineDocument(
            doc_id=doc_id,
            machine_id=machine_id,
            file_path=file_path,
            file_type=file_type,
            content=content,
            upload_date=datetime.now()
        )

class EnhancedLLMProvider:
    def __init__(self):
        self.current_provider = "auto"  # auto, groq, gemini, local
        self.init_providers()
    
    def set_provider(self, provider: str):
        """Set specific LLM provider"""
        valid_providers = ["auto", "groq", "gemini", "local"]
        if provider.lower() in valid_providers:
            self.current_provider = provider.lower()
            logger.info(f"LLM provider set to: {self.current_provider}")
        else:
            logger.error(f"Invalid provider. Choose from: {valid_providers}")
    
    def get_current_provider(self):
        """Get current LLM provider"""
        return self.current_provider
    
    def init_providers(self):
        """Initialize free LLM providers only"""
        # Groq (Free tier available)
        self.groq_configured = (LLMConfig.GROQ_API_KEY and 
                              LLMConfig.GROQ_API_KEY != "your-groq-api-key")
        
        # Gemini (Free tier available)
        self.gemini_configured = (LLMConfig.GEMINI_API_KEY and 
                                LLMConfig.GEMINI_API_KEY != "your-gemini-api-key")
        
        logger.info(f"Free providers configured: "
                   f"Groq={self.groq_configured}, "
                   f"Gemini={self.gemini_configured}")
    
    async def call_groq(self, prompt: str) -> str:
        """Call Groq API (free tier available)"""
        try:
            if not self.groq_configured:
                return "Groq API not configured"
            
            headers = {
                "Authorization": f"Bearer {LLMConfig.GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            
            # Try different free models
            models = [
                "llama3-8b-8192",    # Primary model
                "mixtral-8x7b-32768", # Alternative model
                "gemma-7b-it"        # Another option
            ]
            
            for model in models:
                try:
                    payload = {
                        "model": model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": 1024,
                        "temperature": 0.7
                    }
                    
                    async with aiohttp.ClientSession() as session:
                        async with session.post(
                            "https://api.groq.com/openai/v1/chat/completions",
                            headers=headers,
                            json=payload,
                            timeout=aiohttp.ClientTimeout(total=30)
                        ) as response:
                            if response.status == 200:
                                data = await response.json()
                                return data["choices"][0]["message"]["content"].strip()
                            elif response.status == 429:
                                # Rate limit, try next model
                                continue
                            else:
                                # Try next model on error
                                continue
                                
                except Exception as e:
                    logger.warning(f"Groq model {model} failed: {e}")
                    continue
            
            return "Groq API: All models unavailable"
                        
        except Exception as e:
            logger.error(f"Groq API error: {e}")
            return f"Error calling Groq: {e}"
    
    async def call_gemini(self, prompt: str) -> str:
        """Call Gemini API (free tier available)"""
        try:
            if not self.gemini_configured:
                return "Gemini API not configured"
            
            # Try different Gemini models
            models = [
                "gemini-pro",           # Primary model
                "gemini-1.5-flash"      # Alternative model
            ]
            
            for model in models:
                try:
                    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={LLMConfig.GEMINI_API_KEY}"
                    
                    payload = {
                        "contents": [{
                            "parts": [{"text": prompt}]
                        }],
                        "generationConfig": {
                            "temperature": 0.7,
                            "maxOutputTokens": 1024
                        }
                    }
                    
                    async with aiohttp.ClientSession() as session:
                        async with session.post(
                            url,
                            json=payload,
                            timeout=aiohttp.ClientTimeout(total=30)
                        ) as response:
                            if response.status == 200:
                                data = await response.json()
                                if 'candidates' in data and len(data['candidates']) > 0:
                                    content = data['candidates'][0]['content']['parts'][0]['text']
                                    return content.strip()
                                return "No response generated"
                            elif response.status == 429:
                                # Rate limit, try next model
                                continue
                            else:
                                # Try next model on error
                                continue
                                
                except Exception as e:
                    logger.warning(f"Gemini model {model} failed: {e}")
                    continue
            
            return "Gemini API: All models unavailable"
                        
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            return f"Error calling Gemini: {e}"
    
    async def call_local_fallback(self, prompt: str) -> str:
        """Enhanced local fallback with better maintenance knowledge"""
        maintenance_keywords = {
            'bearing': {
                'response': "BEARING ANALYSIS: Common bearing failure causes include inadequate lubrication (40%), misalignment (25%), contamination (15%), overloading (10%), and fatigue (10%). Recommended actions: 1) Check lubrication levels and quality, 2) Inspect shaft alignment, 3) Look for contamination sources, 4) Verify load conditions, 5) Replace if excessive wear detected.",
                'symptoms': ['noise', 'vibration', 'heat', 'play']
            },
            'pump': {
                'response': "PUMP ANALYSIS: Common pump issues include cavitation (suction problems), impeller wear, mechanical seal failure, and misalignment. Check: 1) Suction pressure and NPSH, 2) Impeller condition and clearances, 3) Mechanical seal integrity, 4) Shaft alignment, 5) Motor coupling condition.",
                'symptoms': ['noise', 'vibration', 'leak', 'flow', 'pressure']
            },
            'motor': {
                'response': "MOTOR ANALYSIS: Common motor problems include overheating, bearing failure, insulation breakdown, and rotor issues. Check: 1) Operating temperature, 2) Vibration levels, 3) Electrical connections and insulation, 4) Current draw and power factor, 5) Bearing condition.",
                'symptoms': ['heat', 'noise', 'vibration', 'current', 'speed']
            }
        }
        
        prompt_lower = prompt.lower()
        
        # Find the most relevant keyword
        best_match = None
        max_score = 0
        
        for keyword, data in maintenance_keywords.items():
            score = 0
            if keyword in prompt_lower:
                score += 10
            
            # Check for related symptoms
            if 'symptoms' in data:
                for symptom in data['symptoms']:
                    if symptom in prompt_lower:
                        score += 5
            
            if score > max_score:
                max_score = score
                best_match = data['response']
        
        if best_match:
            return f"MAINTENANCE ANALYSIS (Local Knowledge Base): {best_match}"
        
        # General maintenance advice if no specific match
        if any(word in prompt_lower for word in ['problem', 'issue', 'failure', 'fault', 'broken']):
            return """GENERAL MAINTENANCE ANALYSIS (Local Knowledge Base): 
            For effective troubleshooting: 1) Document all symptoms and operating conditions, 2) Check recent maintenance history, 3) Verify proper operating procedures, 4) Inspect for obvious physical damage, 5) Measure key parameters (temperature, pressure, vibration), 6) Consult equipment manuals, 7) Consider environmental factors. Always prioritize safety and follow lockout/tagout procedures."""
        
        return "MAINTENANCE ANALYSIS (Local Knowledge Base): Please provide more specific information about the equipment type, symptoms observed, and operating conditions for detailed analysis."
    
    async def get_best_response(self, prompt: str) -> str:
        """Get response from specified provider or try providers in order"""
        # Enhanced prompt for better maintenance responses
        enhanced_prompt = f"""
        You are an expert industrial maintenance engineer with 20+ years of experience in troubleshooting mechanical, electrical, and process equipment. 
        
        User Query: {prompt}
        
        Please provide a detailed, technical response that includes:
        1. Root cause analysis
        2. Systematic troubleshooting approach
        3. Safety considerations
        4. Specific recommendations
        5. Preventive measures
        
        Focus on practical, actionable advice based on industry best practices.
        """
        
        if self.current_provider == "groq" and self.groq_configured:
            response = await self.call_groq(enhanced_prompt)
            if not any(error in response.lower() for error in ['error', 'unavailable', 'not configured']):
                return response
        elif self.current_provider == "gemini" and self.gemini_configured:
            response = await self.call_gemini(enhanced_prompt)
            if not any(error in response.lower() for error in ['error', 'unavailable', 'not configured']):
                return response
        elif self.current_provider == "local":
            return await self.call_local_fallback(enhanced_prompt)
        elif self.current_provider == "auto":
            # Try providers in order of preference
            providers = [
                ("Groq", self.call_groq) if self.groq_configured else None,
                ("Gemini", self.call_gemini) if self.gemini_configured else None,
                ("Local Fallback", self.call_local_fallback)
            ]
            
            # Filter out None providers
            providers = [p for p in providers if p is not None]
            
            for provider_name, provider_func in providers:
                try:
                    logger.info(f"Trying {provider_name} provider...")
                    response = await provider_func(enhanced_prompt)
                    
                    # Check if response indicates an error
                    if not any(error_term in response.lower() for error_term in 
                              ['error', 'quota', 'exceeded', 'not configured', 'unavailable']):
                        logger.info(f"Success with {provider_name}")
                        return response
                    else:
                        logger.warning(f"{provider_name} failed: {response[:100]}...")
                        
                except Exception as e:
                    logger.error(f"{provider_name} provider failed: {e}")
                    continue
            
            # If all providers fail, return local fallback
            return await self.call_local_fallback(enhanced_prompt)
        else:
            return "Invalid provider configuration"

class EnhancedSmartMaintenanceSystem:
    """Enhanced Smart Maintenance Diagnostic System with Flask API"""
    
    def __init__(self):
        self.llm_provider = EnhancedLLMProvider()
        self.document_processor = DocumentProcessor()
        self.equipment_database = {}
        self.diagnosis_history = []
        self.active_sessions = {}
        self.documents = {}
        
        # Initialize Flask app
        self.app = Flask(__name__)
        CORS(self.app)  # Enable CORS for browser requests
        
        # Configure upload settings
        self.app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
        self.app.config['UPLOAD_FOLDER'] = 'uploads'
        
        # Ensure upload directory exists
        os.makedirs(self.app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Thread pool for async operations
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        # Setup Flask routes
        self.setup_routes()
    
    def setup_routes(self):
        """Setup Flask API routes"""
        
        # Health check endpoint
        @self.app.route('/api/health', methods=['GET'])
        def health_check():
            return jsonify({
                'status': 'healthy',
                'timestamp': datetime.now().isoformat(),
                'providers': {
                    'groq': self.llm_provider.groq_configured,
                    'gemini': self.llm_provider.gemini_configured,
                    'local': True
                }
            })
        
        # Equipment management endpoints
        @self.app.route('/api/equipment', methods=['GET'])
        def get_equipment():
            """Get all equipment"""
            equipment_list = []
            for machine_id, info in self.equipment_database.items():
                equipment_list.append({
                    'id': machine_id,
                    'name': info['name'],
                    'type': info['type'],
                    'added_date': info['added_date'].isoformat()
                })
            return jsonify({'equipment': equipment_list})
        
        @self.app.route('/api/equipment', methods=['POST'])
        def add_equipment():
            """Add new equipment"""
            try:
                data = request.json
                machine_id = data.get('id')
                machine_name = data.get('name')
                machine_type = data.get('type')
                
                if not all([machine_id, machine_name, machine_type]):
                    return jsonify({'error': 'Missing required fields: id, name, type'}), 400
                
                if machine_id in self.equipment_database:
                    return jsonify({'error': 'Equipment with this ID already exists'}), 400
                
                self.equipment_database[machine_id] = {
                    'name': machine_name,
                    'type': machine_type,
                    'added_date': datetime.now()
                }
                
                return jsonify({
                    'message': 'Equipment added successfully',
                    'equipment': {
                        'id': machine_id,
                        'name': machine_name,
                        'type': machine_type,
                        'added_date': datetime.now().isoformat()
                    }
                }), 201
                
            except Exception as e:
                return jsonify({'error': str(e)}), 500
        
        # Diagnosis endpoints
        @self.app.route('/api/diagnose', methods=['POST'])
        def diagnose_equipment():
            """Main diagnosis endpoint"""
            try:
                data = request.json
                message = data.get('message')
                equipment_id = data.get('equipment_id')
                session_id = data.get('session_id')
                
                if not message:
                    return jsonify({'error': 'Message is required'}), 400
                
                # Run async diagnosis in thread pool
                future = self.executor.submit(
                    asyncio.run,
                    self.ai_diagnosis_function(message, equipment_id, session_id)
                )
                
                response = future.result(timeout=60)  # 60 second timeout
                
                return jsonify(response)
                
            except Exception as e:
                logger.error(f"Diagnosis error: {e}")
                return jsonify({'error': str(e)}), 500
        
        @self.app.route('/api/sessions', methods=['POST'])
        def create_session():
            """Create new diagnosis session"""
            try:
                data = request.json
                equipment_id = data.get('equipment_id')
                
                if not equipment_id:
                    return jsonify({'error': 'Equipment ID is required'}), 400
                
                if equipment_id not in self.equipment_database:
                    return jsonify({'error': 'Equipment not found'}), 404
                
                session_id = str(uuid.uuid4())
                session = DiagnosisSession(
                    session_id=session_id,
                    machine_id=equipment_id,
                    start_time=datetime.now(),
                    current_diagnosis=None,
                    conversation_history=[],
                    uploaded_files=[],
                    status="active"
                )
                
                self.active_sessions[session_id] = session
                
                return jsonify({
                    'session_id': session_id,
                    'equipment_id': equipment_id,
                    'status': 'active',
                    'created_at': datetime.now().isoformat()
                }), 201
                
            except Exception as e:
                return jsonify({'error': str(e)}), 500
        
        @self.app.route('/api/sessions/<session_id>', methods=['GET'])
        def get_session(session_id):
            """Get session details"""
            if session_id not in self.active_sessions:
                return jsonify({'error': 'Session not found'}), 404
            
            session = self.active_sessions[session_id]
            return jsonify({
                'session_id': session.session_id,
                'equipment_id': session.machine_id,
                'status': session.status,
                'start_time': session.start_time.isoformat(),
                'conversation_count': len(session.conversation_history),
                'uploaded_files': len(session.uploaded_files)
            })
        
        # File upload endpoint
        @self.app.route('/api/upload', methods=['POST'])
        def upload_file():
            """Upload document for equipment"""
            try:
                if 'file' not in request.files:
                    return jsonify({'error': 'No file provided'}), 400
                
                file = request.files['file']
                equipment_id = request.form.get('equipment_id')
                
                if file.filename == '':
                    return jsonify({'error': 'No file selected'}), 400
                
                if not equipment_id:
                    return jsonify({'error': 'Equipment ID is required'}), 400
                
                if equipment_id not in self.equipment_database:
                    return jsonify({'error': 'Equipment not found'}), 404
                
                # Secure filename and save
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                secure_name = f"{equipment_id}_{timestamp}_{filename}"
                file_path = os.path.join(self.app.config['UPLOAD_FOLDER'], secure_name)
                
                file.save(file_path)
                
                # Process document
                doc = self.document_processor.process_document(file_path, equipment_id)
                self.documents[doc.doc_id] = doc
                
                return jsonify({
                    'document_id': doc.doc_id,
                    'filename': filename,
                    'equipment_id': equipment_id,
                    'content_length': len(doc.content),
                    'upload_date': doc.upload_date.isoformat()
                }), 201
                
            except Exception as e:
                return jsonify({'error': str(e)}), 500
        
        # Provider management endpoints
        @self.app.route('/api/providers', methods=['GET'])
        def get_providers():
            """Get available LLM providers and their status"""
            return jsonify({
                'current_provider': self.llm_provider.get_current_provider(),
                'providers': {
                    'groq': {
                        'configured': self.llm_provider.groq_configured,
                        'description': 'Groq API (Free tier available)'
                    },
                    'gemini': {
                        'configured': self.llm_provider.gemini_configured,
                        'description': 'Google Gemini API (Free tier available)'
                    },
                    'local': {
                        'configured': True,
                        'description': 'Local fallback knowledge base'
                    },
                    'auto': {
                        'configured': True,
                        'description': 'Automatic provider selection'
                    }
                }
            })
        
        @self.app.route('/api/providers', methods=['POST'])
        def set_provider():
            """Set LLM provider"""
            try:
                data = request.json
                provider = data.get('provider')
                
                if not provider:
                    return jsonify({'error': 'Provider is required'}), 400
                
                self.llm_provider.set_provider(provider)
                
                return jsonify({
                    'message': f'Provider set to {provider}',
                    'current_provider': self.llm_provider.get_current_provider()
                })
                
            except Exception as e:
                return jsonify({'error': str(e)}), 500
        
        # Test endpoint
        @self.app.route('/api/test', methods=['POST'])
        def test_connection():
            """Test LLM provider connections"""
            try:
                future = self.executor.submit(
                    asyncio.run,
                    self.llm_provider.get_best_response("What are common signs of bearing failure? Respond briefly.")
                )
                
                response = future.result(timeout=30)
                
                return jsonify({
                    'status': 'success',
                    'provider': self.llm_provider.get_current_provider(),
                    'response': response[:200] + "..." if len(response) > 200 else response
                })
                
            except Exception as e:
                return jsonify({'error': str(e)}), 500
        
        @self.app.route('/')
        def index():
            """Serve the main HTML interface"""
            return send_from_directory('.', 'newupdatedui.html')

        @self.app.route('/<path:filename>')
        def serve_static(filename):
            """Serve static files"""
            return send_from_directory('.', filename)
    
    async def ai_diagnosis_function(self, message: str, equipment_id: str = None, session_id: str = None):
        """AI diagnosis function for API"""
        try:
            # Build context if equipment_id is provided
            context = ""
            if equipment_id and equipment_id in self.equipment_database:
                equipment = self.equipment_database[equipment_id]
                context = f"Equipment: {equipment['name']} (Type: {equipment['type']}, ID: {equipment_id})\n\n"
            
            # Add session context if available
            if session_id and session_id in self.active_sessions:
                session = self.active_sessions[session_id]
                if session.conversation_history:
                    context += "Previous conversation:\n"
                    for conv in session.conversation_history[-3:]:  # Last 3 exchanges
                        context += f"User: {conv['user']}\nAI: {conv['ai'][:100]}...\n"
                    context += "\n"
            
            # Get AI response
            full_prompt = context + "Current question: " + message
            response = await self.llm_provider.get_best_response(full_prompt)
            
            # Update session if provided
            if session_id and session_id in self.active_sessions:
                session = self.active_sessions[session_id]
                session.conversation_history.append({
                    'user': message,
                    'ai': response,
                    'timestamp': datetime.now().isoformat()
                })
            
            return {
                'response': response,
                'equipment_id': equipment_id,
                'session_id': session_id,
                'provider': self.llm_provider.get_current_provider(),
                'timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"AI diagnosis error: {e}")
            raise e
    
    def run_api_server(self, host='127.0.0.1', port=5000, debug=False):
        """Run the Flask API server"""
        logger.info(f"Starting Smart Maintenance API server on {host}:{port}")
        logger.info("API Endpoints:")
        logger.info("  GET  /api/health - Health check")
        logger.info("  GET  /api/equipment - Get all equipment")
        logger.info("  POST /api/equipment - Add equipment")
        logger.info("  POST /api/diagnose - Diagnose equipment")
        logger.info("  POST /api/sessions - Create diagnosis session")
        logger.info("  GET  /api/sessions/<id> - Get session details")
        logger.info("  POST /api/upload - Upload document")
        logger.info("  GET  /api/providers - Get provider status")
        logger.info("  POST /api/providers - Set provider")
        logger.info("  POST /api/test - Test connection")
        
        self.app.run(host=host, port=port, debug=debug, threaded=True)
    
    # Keep existing interactive methods for backward compatibility
    def display_menu(self):
        """Display main menu"""
        print("\n" + "="*50)
        print("Smart Maintenance Diagnostic System")
        print("="*50)
        print("1. Add Equipment")
        print("2. Start Diagnosis Session")
        print("3. Upload Documents")
        print("4. View Equipment List")
        print("5. Test LLM Connection")
        print("6. Change LLM Provider")
        print("7. Chat with AI")
        print("8. Start API Server")
        print("9. Exit")
        print("="*50)
    
    async def add_equipment_interactive(self):
        """Interactive method to add equipment"""
        print("\n--- Add New Equipment ---")
        
        while True:
            machine_id = input("Enter machine ID: ").strip()
            if machine_id:
                if machine_id in self.equipment_database:
                    print("❌ Equipment with this ID already exists!")
                    continue
                break
            print("❌ Machine ID cannot be empty!")
        
        while True:
            machine_name = input("Enter machine name: ").strip()
            if machine_name:
                break
            print("❌ Machine name cannot be empty!")
        
        while True:
            machine_type = input("Enter machine type (e.g., pump, motor, compressor): ").strip()
            if machine_type:
                break
            print("❌ Machine type cannot be empty!")
        
        # Add to database
        self.equipment_database[machine_id] = {
            'name': machine_name,
            'type': machine_type,
            'added_date': datetime.now()
        }
        
        print(f"✅ Equipment '{machine_name}' (ID: {machine_id}) added successfully!")
    
    def view_equipment_list(self):
        """Display all equipment in the database"""
        if not self.equipment_database:
            print("\n❌ No equipment found in the database.")
            return
        
        print("\n--- Equipment List ---")
        print(f"{'ID':<15} {'Name':<25} {'Type':<20} {'Added Date':<20}")
        print("-" * 80)
        
        for machine_id, info in self.equipment_database.items():
            print(f"{machine_id:<15} {info['name']:<25} {info['type']:<20} {info['added_date'].strftime('%Y-%m-%d %H:%M'):<20}")
    
    async def start_diagnosis_session_interactive(self):
        """Interactive diagnosis session"""
        if not self.equipment_database:
            print("\n❌ No equipment found. Please add equipment first.")
            return
        
        print("\n--- Start Diagnosis Session ---")
        self.view_equipment_list()
        
        while True:
            machine_id = input("\nEnter machine ID to diagnose: ").strip()
            if machine_id in self.equipment_database:
                break
            print("❌ Invalid machine ID. Please try again.")
        
        # Create session
        session_id = str(uuid.uuid4())
        session = DiagnosisSession(
            session_id=session_id,
            machine_id=machine_id,
            start_time=datetime.now(),
            current_diagnosis=None,
            conversation_history=[],
            uploaded_files=[],
            status="active"
        )
        
        self.active_sessions[session_id] = session
        equipment = self.equipment_database[machine_id]
        
        print(f"\n🔧 Starting diagnosis session for: {equipment['name']} ({equipment['type']})")
        print(f"Session ID: {session_id}")
        print("\nDescribe the problem or ask questions. Type 'exit' to end session.\n")
        
        while True:
            user_input = input("You: ").strip()
            
            if user_input.lower() in ['exit', 'quit', 'end']:
                session.status = "completed"
                print("✅ Diagnosis session ended.")
                break
            
            if not user_input:
                continue
            
            try:
                print("🤖 AI is analyzing... Please wait.")
                
                # Get AI response
                response_data = await self.ai_diagnosis_function(user_input, machine_id, session_id)
                response = response_data['response']
                
                print(f"\nAI Expert: {response}\n")
                
            except Exception as e:
                print(f"❌ Error getting AI response: {e}")
    
    async def upload_documents_interactive(self):
        """Interactive document upload"""
        if not self.equipment_database:
            print("\n❌ No equipment found. Please add equipment first.")
            return
        
        print("\n--- Upload Documents ---")
        self.view_equipment_list()
        
        while True:
            machine_id = input("\nEnter machine ID for document upload: ").strip()
            if machine_id in self.equipment_database:
                break
            print("❌ Invalid machine ID. Please try again.")
        
        while True:
            file_path = input("Enter file path (PDF, DOCX, TXT): ").strip().strip('"')
            
            if not file_path:
                print("❌ File path cannot be empty!")
                continue
            
            if not os.path.exists(file_path):
                print("❌ File not found!")
                continue
            
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in ['.pdf', '.docx', '.txt', '.md']:
                print("❌ Unsupported file type. Please use PDF, DOCX, or TXT files.")
                continue
            
            break
        
        try:
            print("📄 Processing document...")
            doc = self.document_processor.process_document(file_path, machine_id)
            self.documents[doc.doc_id] = doc
            
            print(f"✅ Document uploaded successfully!")
            print(f"Document ID: {doc.doc_id}")
            print(f"Content length: {len(doc.content)} characters")
            
        except Exception as e:
            print(f"❌ Error uploading document: {e}")
    
    async def test_llm_connection_interactive(self):
        """Test LLM provider connection"""
        print("\n--- Testing LLM Connection ---")
        print(f"Current provider: {self.llm_provider.get_current_provider()}")
        
        test_prompt = "What are common signs of bearing failure? Respond briefly."
        
        try:
            print("🔄 Testing connection...")
            response = await self.llm_provider.get_best_response(test_prompt)
            
            print("✅ Connection successful!")
            print(f"Provider: {self.llm_provider.get_current_provider()}")
            print(f"Response: {response[:200]}{'...' if len(response) > 200 else ''}")
            
        except Exception as e:
            print(f"❌ Connection failed: {e}")
    
    def change_llm_provider_interactive(self):
        """Interactive LLM provider selection"""
        print("\n--- Change LLM Provider ---")
        print("Available providers:")
        print("1. Auto (automatic selection)")
        print("2. Groq (free tier)")
        print("3. Gemini (free tier)")
        print("4. Local (fallback)")
        
        print(f"\nCurrent provider: {self.llm_provider.get_current_provider()}")
        print(f"Groq configured: {self.llm_provider.groq_configured}")
        print(f"Gemini configured: {self.llm_provider.gemini_configured}")
        
        while True:
            choice = input("\nSelect provider (1-4): ").strip()
            
            if choice == "1":
                self.llm_provider.set_provider("auto")
                break
            elif choice == "2":
                self.llm_provider.set_provider("groq")
                break
            elif choice == "3":
                self.llm_provider.set_provider("gemini")
                break
            elif choice == "4":
                self.llm_provider.set_provider("local")
                break
            else:
                print("❌ Invalid choice. Please enter 1-4.")
        
        print(f"✅ Provider changed to: {self.llm_provider.get_current_provider()}")
    
    async def chat_with_ai_interactive(self):
        """Interactive AI chat"""
        print("\n--- Chat with AI Expert ---")
        print("Ask any maintenance-related questions. Type 'exit' to return to main menu.\n")
        
        while True:
            user_input = input("You: ").strip()
            
            if user_input.lower() in ['exit', 'quit', 'back']:
                break
            
            if not user_input:
                continue
            
            try:
                print("🤖 AI is thinking... Please wait.")
                response = await self.llm_provider.get_best_response(user_input)
                print(f"\nAI Expert: {response}\n")
                
            except Exception as e:
                print(f"❌ Error: {e}\n")
    
    def start_api_server_interactive(self):
        """Interactive API server startup"""
        print("\n--- Start API Server ---")
        
        host = input("Enter host (default: 127.0.0.1): ").strip() or "127.0.0.1"
        
        while True:
            port_input = input("Enter port (default: 5000): ").strip() or "5000"
            try:
                port = int(port_input)
                if 1 <= port <= 65535:
                    break
                else:
                    print("❌ Port must be between 1 and 65535.")
            except ValueError:
                print("❌ Invalid port number.")
        
        debug_input = input("Enable debug mode? (y/N): ").strip().lower()
        debug = debug_input in ['y', 'yes']
        
        print(f"\n🚀 Starting API server on {host}:{port}")
        print("Press Ctrl+C to stop the server.")
        
        try:
            self.run_api_server(host=host, port=port, debug=debug)
        except KeyboardInterrupt:
            print("\n✅ Server stopped.")
        except Exception as e:
            print(f"❌ Server error: {e}")
    
    async def run_interactive(self):
        """Main interactive loop"""
        print("🔧 Smart Maintenance Diagnostic System")
        print("Initializing...")
        
        # Test initial provider setup
        print(f"LLM Provider: {self.llm_provider.get_current_provider()}")
        print(f"Groq configured: {self.llm_provider.groq_configured}")
        print(f"Gemini configured: {self.llm_provider.gemini_configured}")
        
        while True:
            self.display_menu()
            
            choice = input("\nEnter your choice (1-9): ").strip()
            
            try:
                if choice == "1":
                    await self.add_equipment_interactive()
                elif choice == "2":
                    await self.start_diagnosis_session_interactive()
                elif choice == "3":
                    await self.upload_documents_interactive()
                elif choice == "4":
                    self.view_equipment_list()
                elif choice == "5":
                    await self.test_llm_connection_interactive()
                elif choice == "6":
                    self.change_llm_provider_interactive()
                elif choice == "7":
                    await self.chat_with_ai_interactive()
                elif choice == "8":
                    self.start_api_server_interactive()
                elif choice == "9":
                    print("👋 Goodbye!")
                    break
                else:
                    print("❌ Invalid choice. Please enter 1-9.")
            
            except KeyboardInterrupt:
                print("\n\n👋 Goodbye!")
                break
            except Exception as e:
                print(f"❌ Error: {e}")
                logger.error(f"Interactive error: {e}")

# Main execution
async def main():
    """Main function"""
    system = EnhancedSmartMaintenanceSystem()
    
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "api":
        # Start API server directly
        host = sys.argv[2] if len(sys.argv) > 2 else "127.0.0.1"
        port = int(sys.argv[3]) if len(sys.argv) > 3 else 5000
        debug = len(sys.argv) > 4 and sys.argv[4].lower() == "debug"
        
        system.run_api_server(host=host, port=port, debug=debug)
    else:
        # Start interactive mode
        await system.run_interactive()

if __name__ == "__main__":
    asyncio.run(main())