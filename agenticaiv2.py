import asyncio
import json
import logging
import sqlite3
import os
import cv2
import numpy as np
import base64
import mimetypes
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from enum import Enum

import google.generativeai as genai
import requests
import aiohttp
from dotenv import load_dotenv
import PyPDF2
import docx
from pathlib import Path
import hashlib

load_dotenv()

class LLMConfig:
    """Configuration for LLM providers"""
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "your-gemini-api-key")

gemini_key = os.getenv("GEMINI_API_KEY")
print(f"Gemini API Key status: {'Configured' if gemini_key else 'Not configured'}")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enhanced Data Models
@dataclass
class SensorData:
    sensor_id: str
    timestamp: datetime
    temperature: float
    vibration: float
    pressure: float
    humidity: float
    sound_level: float

@dataclass
class MachineDocument:
    doc_id: str
    machine_id: str
    file_path: str
    file_type: str
    content: str
    upload_date: datetime
    summary: str = ""

@dataclass
class VisualInspection:
    inspection_id: str
    machine_id: str
    file_path: str
    file_type: str  # image or video
    defects_detected: List[Dict]
    timestamp: datetime
    ai_analysis: str = ""

@dataclass
class FaultDiagnosis:
    fault_id: str
    machine_id: str
    severity: str
    confidence: float
    description: str
    root_cause: str
    recommended_actions: List[str]
    estimated_downtime: int
    supporting_evidence: List[str]  # References to documents/images
    conversation_history: Optional[List[Dict]] = None

@dataclass
class DiagnosisSession:
    session_id: str
    machine_id: str
    start_time: datetime
    current_diagnosis: Optional[FaultDiagnosis]
    conversation_history: List[Dict]
    uploaded_files: List[str]
    status: str  # active, resolved, escalated

class MaintenanceType(Enum):
    PREVENTIVE = "preventive"
    CORRECTIVE = "corrective"
    PREDICTIVE = "predictive"
    EMERGENCY = "emergency"

@dataclass
class MaintenanceTask:
    task_id: str
    equipment_id: str
    task_type: MaintenanceType
    priority: int
    description: str
    scheduled_date: datetime
    estimated_duration: int
    required_parts: List[str]
    assigned_technician: Optional[str] = None

# Enhanced Document Processing
class DocumentProcessor:
    def __init__(self, storage_dir: str = "uploads"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
    def save_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded file and return new path"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        
        # Copy file to storage directory
        import shutil
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def process_document(self, file_path: str, machine_id: str) -> MachineDocument:
        """Process uploaded document and extract content"""
        file_path_obj = Path(file_path)
        file_type = file_path_obj.suffix.lower()
        
        # Extract text based on file type
        if file_type == '.pdf':
            content = self.extract_text_from_pdf(file_path)
        elif file_type == '.docx':
            content = self.extract_text_from_docx(file_path)
        elif file_type in ['.txt', '.md']:
            content = self.extract_text_from_txt(file_path)
        else:
            content = f"Unsupported file type: {file_type}"
        
        # Generate document ID
        doc_id = hashlib.md5(f"{machine_id}_{file_path_obj.name}_{datetime.now()}".encode()).hexdigest()[:12]
        
        return MachineDocument(
            doc_id=doc_id,
            machine_id=machine_id,
            file_path=file_path,
            file_type=file_type,
            content=content,
            upload_date=datetime.now()
        )

# Enhanced Visual Processing
class EnhancedVisionProcessor:
    def __init__(self, llm_provider, storage_dir: str = "visual_uploads"):
        self.llm_provider = llm_provider
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
    
    def save_visual_file(self, file_path: str, machine_id: str) -> str:
        """Save uploaded visual file"""
        original_path = Path(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{machine_id}_{timestamp}_{original_path.name}"
        new_path = self.storage_dir / new_filename
        
        import shutil
        shutil.copy2(file_path, new_path)
        return str(new_path)
    
    def analyze_image(self, image_path: str) -> Dict[str, Any]:
        """Enhanced image analysis with OpenCV"""
        try:
            image = cv2.imread(image_path)
            if image is None:
                return {"error": "Could not load image"}
            
            # Multiple analysis techniques
            analysis_results = {}
            
            # 1. Basic defect detection
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            edges = cv2.Canny(gray, 50, 150)
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            defects = []
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > 500:  # Threshold for significant defects
                    x, y, w, h = cv2.boundingRect(contour)
                    defects.append({
                        'location': {'x': int(x), 'y': int(y), 'width': int(w), 'height': int(h)},
                        'area': float(area),
                        'type': 'potential_defect'
                    })
            
            # 2. Corrosion detection (rust-like colors)
            hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
            rust_lower = np.array([0, 50, 50])
            rust_upper = np.array([20, 255, 255])
            rust_mask = cv2.inRange(hsv, rust_lower, rust_upper)
            rust_area = cv2.countNonZero(rust_mask)
            
            # 3. Crack detection using morphological operations
            kernel = np.ones((3,3), np.uint8)
            morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
            crack_edges = cv2.Canny(morph, 100, 200)
            crack_lines = cv2.HoughLinesP(crack_edges, 1, np.pi/180, threshold=50, minLineLength=30, maxLineGap=10)
            
            analysis_results = {
                'defects_found': len(defects),
                'defects': defects,
                'rust_area_pixels': int(rust_area),
                'potential_cracks': len(crack_lines) if crack_lines is not None else 0,
                'image_dimensions': {'width': image.shape[1], 'height': image.shape[0]},
                'analysis_timestamp': datetime.now().isoformat()
            }
            
            return analysis_results
            
        except Exception as e:
            logger.error(f"Vision processing error: {e}")
            return {"error": str(e)}
    
    def analyze_video(self, video_path: str) -> Dict[str, Any]:
        """Basic video analysis for defect detection"""
        try:
            cap = cv2.VideoCapture(video_path)
            if not cap.isOpened():
                return {"error": "Could not open video"}
            
            frame_count = 0
            total_defects = 0
            frame_analyses = []
            
            # Analyze every 30th frame to avoid processing every frame
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                
                if frame_count % 30 == 0:  # Process every 30th frame
                    # Save frame temporarily
                    temp_frame_path = f"temp_frame_{frame_count}.jpg"
                    cv2.imwrite(temp_frame_path, frame)
                    
                    # Analyze frame
                    frame_analysis = self.analyze_image(temp_frame_path)
                    if 'defects_found' in frame_analysis:
                        total_defects += frame_analysis['defects_found']
                        frame_analyses.append({
                            'frame_number': frame_count,
                            'defects': frame_analysis['defects_found'],
                            'timestamp': frame_count / cap.get(cv2.CAP_PROP_FPS) if cap.get(cv2.CAP_PROP_FPS) > 0 else 0
                        })
                    
                    # Clean up temp file
                    if os.path.exists(temp_frame_path):
                        os.remove(temp_frame_path)
                
                frame_count += 1
            
            cap.release()
            
            return {
                'total_frames': frame_count,
                'frames_analyzed': len(frame_analyses),
                'total_defects_detected': total_defects,
                'frame_analyses': frame_analyses,
                'analysis_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Video processing error: {e}")
            return {"error": str(e)}
    
    async def analyze_with_gemini_vision(self, file_path: str, machine_context: str = "") -> str:
        """Analyze visual content with Gemini Pro Vision"""
        try:
            # Get computer vision analysis first
            file_ext = Path(file_path).suffix.lower()
            is_video = file_ext in ['.mp4', '.avi', '.mov', '.mkv']
            
            if is_video:
                cv_analysis = self.analyze_video(file_path)
            else:
                cv_analysis = self.analyze_image(file_path)
            
            prompt = f"""
            As an expert industrial maintenance engineer, analyze the following computer vision inspection results:
            
            Machine Context: {machine_context}
            File Type: {"Video" if is_video else "Image"}
            File Path: {file_path}
            
            COMPUTER VISION ANALYSIS RESULTS:
            {json.dumps(cv_analysis, indent=2)}
            
            Based on these computer vision findings, please provide:
            1. Interpretation of detected defects and their significance
            2. Assessment of rust/corrosion areas (if any)
            3. Evaluation of potential cracks or structural issues
            4. Overall condition assessment
            5. Immediate safety concerns (if any)
            6. Recommended maintenance actions
            7. Priority level for addressing identified issues
            
            Provide specific, actionable recommendations based on the computer vision analysis results.
            Consider the machine context and typical failure modes for this type of equipment.
            """
            
            response = await self.llm_provider.get_best_response(prompt)
            return response
            
        except Exception as e:
            logger.error(f"Gemini vision analysis error: {e}")
            return f"Visual analysis error: {e}"

# Enhanced LLM Provider - Reset to use Gemini
class EnhancedLLMProvider:
    def __init__(self):
        self.init_gemini()
    
    def init_gemini(self):
        """Initialize Gemini provider"""
        try:
            if LLMConfig.GEMINI_API_KEY and LLMConfig.GEMINI_API_KEY != "your-gemini-api-key":
                genai.configure(api_key=LLMConfig.GEMINI_API_KEY)
                self.gemini_model = genai.GenerativeModel('gemini-1.5-flash')
                logger.info("Gemini model initialized successfully")
            else:
                self.gemini_model = None
                logger.warning("Gemini API key not configured")
        except Exception as e:
            logger.warning(f"Error initializing Gemini: {e}")
            self.gemini_model = None
    
    async def call_gemini(self, prompt: str) -> str:
        """Call Gemini API"""
        try:
            if not self.gemini_model:
                return "Gemini API not configured. Please set GEMINI_API_KEY environment variable."
                
            response = self.gemini_model.generate_content(prompt)
            return response.text
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            return f"Error calling Gemini: {e}"
    
    async def get_best_response(self, prompt: str) -> str:
        """Get response from Gemini"""
        return await self.call_gemini(prompt)
    
    async def analyze_with_context(self, prompt: str, context_docs: List[MachineDocument], visual_data: Optional[List[VisualInspection]] = None) -> str:
        """Analyze with machine-specific context"""
        enhanced_prompt = f"""
        MACHINE-SPECIFIC CONTEXT:
        {self._format_document_context(context_docs)}
        
        VISUAL INSPECTION DATA:
        {self._format_visual_context(visual_data) if visual_data else "No visual data available"}
        
        MAINTENANCE EXPERT ANALYSIS REQUEST:
        {prompt}
        
        Please provide analysis based on the specific machine documentation and visual inspection data provided above.
        Reference specific sections of the documentation when relevant.
        """
        
        return await self.get_best_response(enhanced_prompt)
    
    def _format_document_context(self, docs: List[MachineDocument]) -> str:
        """Format document context for LLM"""
        if not docs:
            return "No machine documentation available"
        
        context = ""
        for doc in docs:
            context += f"""
            Document: {doc.file_path}
            Type: {doc.file_type}
            Content Summary: {doc.content[:1000]}...
            ---
            """
        return context
    
    def _format_visual_context(self, visual_data: List[VisualInspection]) -> str:
        """Format visual inspection context"""
        if not visual_data:
            return "No visual inspection data available"
        
        context = ""
        for inspection in visual_data:
            context += f"""
            Visual Inspection: {inspection.file_path}
            Type: {inspection.file_type}
            Defects Detected: {len(inspection.defects_detected)}
            AI Analysis: {inspection.ai_analysis[:500]}...
            ---
            """
        return context
    
    async def close(self):
        """Cleanup resources"""
        pass

# Enhanced Knowledge Base
class EnhancedMaintenanceKnowledgeBase:
    def __init__(self, db_path: str = "enhanced_maintenance.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Initialize enhanced SQLite database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Equipment table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS equipment (
                id TEXT PRIMARY KEY,
                name TEXT,
                type TEXT,
                model TEXT,
                manufacturer TEXT,
                installation_date DATE,
                last_maintenance DATE
            )
        ''')
        
        # Machine documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS machine_documents (
                doc_id TEXT PRIMARY KEY,
                machine_id TEXT,
                file_path TEXT,
                file_type TEXT,
                content TEXT,
                summary TEXT,
                upload_date TIMESTAMP,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Visual inspections table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS visual_inspections (
                inspection_id TEXT PRIMARY KEY,
                machine_id TEXT,
                file_path TEXT,
                file_type TEXT,
                defects_data TEXT,
                ai_analysis TEXT,
                timestamp TIMESTAMP,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Diagnosis sessions table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS diagnosis_sessions (
                session_id TEXT PRIMARY KEY,
                machine_id TEXT,
                start_time TIMESTAMP,
                current_diagnosis TEXT,
                conversation_history TEXT,
                uploaded_files TEXT,
                status TEXT,
                FOREIGN KEY (machine_id) REFERENCES equipment (id)
            )
        ''')
        
        # Enhanced maintenance history
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS maintenance_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                equipment_id TEXT,
                maintenance_date DATE,
                type TEXT,
                description TEXT,
                technician TEXT,
                cost REAL,
                downtime_hours INTEGER,
                supporting_documents TEXT,
                FOREIGN KEY (equipment_id) REFERENCES equipment (id)
            )
        ''')
        
        # Fault patterns table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS fault_patterns (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                equipment_type TEXT,
                symptoms TEXT,
                root_cause TEXT,
                solution TEXT,
                confidence_score REAL,
                supporting_evidence TEXT
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def add_machine_document(self, document: MachineDocument) -> bool:
        """Add machine document to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO machine_documents 
                (doc_id, machine_id, file_path, file_type, content, summary, upload_date)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                document.doc_id,
                document.machine_id,
                document.file_path,
                document.file_type,
                document.content,
                document.summary,
                document.upload_date
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding machine document: {e}")
            return False
    
    def get_machine_documents(self, machine_id: str) -> List[MachineDocument]:
        """Get all documents for a specific machine"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM machine_documents WHERE machine_id = ?
        ''', (machine_id,))
        
        results = cursor.fetchall()
        conn.close()
        
        documents = []
        for row in results:
            doc = MachineDocument(
                doc_id=row[0],
                machine_id=row[1],
                file_path=row[2],
                file_type=row[3],
                content=row[4],
                upload_date=datetime.fromisoformat(row[6]),
                summary=row[5] or ""
            )
            documents.append(doc)
        
        return documents
    
    def add_visual_inspection(self, inspection: VisualInspection) -> bool:
        """Add visual inspection to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO visual_inspections 
                (inspection_id, machine_id, file_path, file_type, defects_data, ai_analysis, timestamp)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                inspection.inspection_id,
                inspection.machine_id,
                inspection.file_path,
                inspection.file_type,
                json.dumps(inspection.defects_detected),
                inspection.ai_analysis,
                inspection.timestamp
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding visual inspection: {e}")
            return False
    
    def get_visual_inspections(self, machine_id: str) -> List[VisualInspection]:
        """Get all visual inspections for a machine"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM visual_inspections WHERE machine_id = ?
        ''', (machine_id,))
        
        results = cursor.fetchall()
        conn.close()
        
        inspections = []
        for row in results:
            inspection = VisualInspection(
                inspection_id=row[0],
                machine_id=row[1],
                file_path=row[2],
                file_type=row[3],
                defects_detected=json.loads(row[4]) if row[4] else [],
                timestamp=datetime.fromisoformat(row[6]),
                ai_analysis=row[5] or ""
            )
            inspections.append(inspection)
        
        return inspections
    
    def save_diagnosis_session(self, session: DiagnosisSession) -> bool:
        """Save diagnosis session"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO diagnosis_sessions 
                (session_id, machine_id, start_time, current_diagnosis, conversation_history, uploaded_files, status)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                session.session_id,
                session.machine_id,
                session.start_time,
                json.dumps(session.current_diagnosis.__dict__ if session.current_diagnosis else {}),
                json.dumps(session.conversation_history),
                json.dumps(session.uploaded_files),
                session.status
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error saving diagnosis session: {e}")
            return False
    
    def get_diagnosis_session(self, session_id: str) -> Optional[DiagnosisSession]:
        """Retrieve diagnosis session"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('SELECT * FROM diagnosis_sessions WHERE session_id = ?', (session_id,))
            result = cursor.fetchone()
            conn.close()
            
            if result:
                diagnosis_data = json.loads(result[3]) if result[3] else {}
                diagnosis = FaultDiagnosis(**diagnosis_data) if diagnosis_data else None
                
                return DiagnosisSession(
                    session_id=result[0],
                    machine_id=result[1],
                    start_time=datetime.fromisoformat(result[2]),
                    current_diagnosis=diagnosis,
                    conversation_history=json.loads(result[4]),
                    uploaded_files=json.loads(result[5]),
                    status=result[6]
                )
            return None
        except Exception as e:
            logger.error(f"Error retrieving diagnosis session: {e}")
            return None
    
    def add_equipment(self, equipment_data: Dict) -> bool:
        """Add new equipment to database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO equipment (id, name, type, model, manufacturer, installation_date, last_maintenance)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                equipment_data['id'],
                equipment_data['name'],
                equipment_data['type'],
                equipment_data.get('model', 'Unknown'),
                equipment_data.get('manufacturer', 'Unknown'),
                equipment_data.get('installation_date', datetime.now().strftime('%Y-%m-%d')),
                equipment_data.get('last_maintenance', 'Never')
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding equipment: {e}")
            return False
    
    def get_all_equipment(self) -> List[Dict]:
        """Get all equipment from database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM equipment')
        equipment = cursor.fetchall()
        columns = [description[0] for description in cursor.description]
        conn.close()
        
        return [dict(zip(columns, row)) for row in equipment]

    def add_maintenance_record(self, maintenance_data: Dict) -> bool:
        """Add maintenance record"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT INTO maintenance_history 
                (equipment_id, maintenance_date, type, description, technician, cost, downtime_hours, supporting_documents)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                maintenance_data['equipment_id'],
                maintenance_data['maintenance_date'],
                maintenance_data['type'],
                maintenance_data['description'],
                maintenance_data.get('technician', 'Unknown'),
                maintenance_data.get('cost', 0.0),
                maintenance_data.get('downtime_hours', 0),
                json.dumps(maintenance_data.get('supporting_documents', []))
            ))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding maintenance record: {e}")
            return False

# Enhanced Diagnostic Engine
class EnhancedDiagnosticEngine:
    def __init__(self, knowledge_base: EnhancedMaintenanceKnowledgeBase, llm_provider: EnhancedLLMProvider):
        self.knowledge_base = knowledge_base
        self.llm_provider = llm_provider
    
    async def comprehensive_diagnosis(self, machine_id: str, sensor_data: Optional[Dict] = None, visual_data: Optional[List[VisualInspection]] = None, user_description: str = "") -> FaultDiagnosis:
        """Perform comprehensive diagnosis using all available data"""
        
        # Get machine documents
        machine_docs = self.knowledge_base.get_machine_documents(machine_id)
        
        # Get visual inspections if not provided
        if visual_data is None:
            visual_data = self.knowledge_base.get_visual_inspections(machine_id)
        
        # Prepare comprehensive analysis prompt
        prompt = f"""
        COMPREHENSIVE MACHINE DIAGNOSIS REQUEST
        
        Machine ID: {machine_id}
        User Description: {user_description}
        
        AVAILABLE DATA:
        1. Machine Documentation: {len(machine_docs)} documents available
        2. Visual Inspections: {len(visual_data)} inspections available
        3. Sensor Data: {"Available" if sensor_data else "Not available"}
        
        ANALYSIS REQUEST:
        Based on the machine-specific documentation and visual inspection data, please provide:
        
        1. PRIMARY DIAGNOSIS:
           - Most likely root cause
           - Confidence level (0.0-1.0)
           - Severity assessment (low/medium/high/critical)
        
        2. SUPPORTING EVIDENCE:
           - Reference specific documentation sections
           - Cite visual inspection findings
           - Correlate with sensor anomalies
        
        3. RECOMMENDED DIAGNOSTIC STEPS:
           - Immediate actions to confirm diagnosis
           - Additional tests or inspections needed
           - Safety precautions
        
        4. ALTERNATIVE DIAGNOSES:
           - Other possible causes
           - How to differentiate between them
        
        5. REPAIR STRATEGY:
           - Step-by-step repair process
           - Required parts and tools
           - Estimated time and complexity
        
        Please provide specific, actionable recommendations based on the machine's actual documentation and inspection data.
        """
        
        # Get AI analysis with context
        ai_response = await self.llm_provider.analyze_with_context(prompt, machine_docs, visual_data)
        
        # Parse response and create diagnosis
        diagnosis = FaultDiagnosis(
            fault_id=f"DIAG_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            machine_id=machine_id,
            severity="medium",  # Default, should be parsed from AI response
            confidence=0.7,     # Default, should be parsed from AI response
            description=f"Gemini-assisted comprehensive diagnosis for {machine_id}",
            root_cause="Analysis in progress",
            recommended_actions=["Review Gemini analysis", "Implement suggested diagnostic steps"],
            estimated_downtime=4,
            supporting_evidence=[doc.file_path for doc in machine_docs] + [vis.file_path for vis in visual_data],
            conversation_history=[{
                "timestamp": datetime.now().isoformat(),
                "type": "diagnosis",
                "content": ai_response,
                "user_input": user_description
            }]
        )
        
        return diagnosis
    
    async def interactive_diagnosis_session(self, machine_id: str) -> DiagnosisSession:
        """Start an interactive diagnosis session"""
        session = DiagnosisSession(
            session_id=f"SESSION_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            machine_id=machine_id,
            start_time=datetime.now(),
            current_diagnosis=None,
            conversation_history=[],
            uploaded_files=[],
            status="active"
        )
        
        # Save initial session
        self.knowledge_base.save_diagnosis_session(session)
        return session
    
    async def continue_diagnosis_conversation(self, session_id: str, user_message: str, uploaded_files: Optional[List[str]] = None) -> str:
        """Continue diagnosis conversation"""
        session = self.knowledge_base.get_diagnosis_session(session_id)
        if not session:
            return "Session not found. Please start a new diagnosis session."
        
        # Process any uploaded files
        if uploaded_files:
            session.uploaded_files.extend(uploaded_files)
        
        # Get machine context
        machine_docs = self.knowledge_base.get_machine_documents(session.machine_id)
        visual_data = self.knowledge_base.get_visual_inspections(session.machine_id)
        
        # Build conversation context
        conversation_context = "\n".join([
            f"[{msg['timestamp']}] {msg['type']}: {msg['content'][:200]}..."
            for msg in session.conversation_history[-5:]  # Last 5 messages
        ])
        
        enhanced_prompt = f"""
        CONTINUING DIAGNOSIS SESSION
        
        Session ID: {session_id}
        Machine ID: {session.machine_id}
        
        CONVERSATION HISTORY:
        {conversation_context}
        
        CURRENT USER MESSAGE:
        {user_message}
        
        RECENTLY UPLOADED FILES:
        {uploaded_files if uploaded_files else "None"}
        
        Please continue the diagnosis conversation, building on previous interactions.
        Reference machine documentation and visual inspection data when relevant.
        Provide specific, actionable guidance based on the conversation context.
        """
        
        # Get AI response
        response = await self.llm_provider.analyze_with_context(enhanced_prompt, machine_docs, visual_data)
        
        # Update conversation history
        session.conversation_history.extend([
            {
                "timestamp": datetime.now().isoformat(),
                "type": "user",
                "content": user_message
            },
            {
                "timestamp": datetime.now().isoformat(),
                "type": "assistant",
                "content": response
            }
        ])
        
        # Save updated session
        self.knowledge_base.save_diagnosis_session(session)
        
        return response

# Enhanced Smart Maintenance System
class EnhancedSmartMaintenanceSystem:
    def __init__(self):
        self.knowledge_base = EnhancedMaintenanceKnowledgeBase()
        self.llm_provider = EnhancedLLMProvider()
        self.diagnostic_engine = EnhancedDiagnosticEngine(self.knowledge_base, self.llm_provider)
        self.document_processor = DocumentProcessor()
        self.vision_processor = EnhancedVisionProcessor(self.llm_provider)
    
    async def upload_machine_document(self, machine_id: str, file_path: str) -> Dict[str, Any]:
        """Upload and process machine document"""
        try:
            # Process document
            document = self.document_processor.process_document(file_path, machine_id)
            
            # Generate summary using Gemini
            summary_prompt = f"""
            Analyze this machine documentation and provide a concise summary:
            
            File: {document.file_path}
            Type: {document.file_type}
            Content: {document.content[:2000]}...
            
            Please provide:
            1. Document type and purpose
            2. Key technical specifications
            3. Important maintenance procedures
            4. Critical safety information
            5. Troubleshooting guidance available
            """
            
            document.summary = await self.llm_provider.get_best_response(summary_prompt)
            
            # Save to knowledge base
            success = self.knowledge_base.add_machine_document(document)
            
            return {
                "success": success,
                "document_id": document.doc_id,
                "summary": document.summary,
                "message": "Document uploaded and processed successfully" if success else "Error processing document"
            }
            
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            return {"success": False, "error": str(e)}
    
    async def upload_visual_inspection(self, machine_id: str, file_path: str) -> Dict[str, Any]:
        """Upload and analyze visual inspection file"""
        try:
            # Determine file type
            file_ext = Path(file_path).suffix.lower()
            is_video = file_ext in ['.mp4', '.avi', '.mov', '.mkv']
            
            # Analyze with computer vision
            if is_video:
                cv_analysis = self.vision_processor.analyze_video(file_path)
            else:
                cv_analysis = self.vision_processor.analyze_image(file_path)
            
            # Get machine context for AI analysis
            machine_docs = self.knowledge_base.get_machine_documents(machine_id)
            machine_context = f"Machine {machine_id} with {len(machine_docs)} documentation files available"
            
            # Analyze with Gemini Vision
            ai_analysis = await self.vision_processor.analyze_with_gemini_vision(file_path, machine_context)
            
            # Create visual inspection record
            inspection = VisualInspection(
                inspection_id=f"INSP_{machine_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                machine_id=machine_id,
                file_path=file_path,
                file_type="video" if is_video else "image",
                defects_detected=cv_analysis.get('defects', []),
                timestamp=datetime.now(),
                ai_analysis=ai_analysis
            )
            
            # Save to knowledge base
            success = self.knowledge_base.add_visual_inspection(inspection)
            
            return {
                "success": success,
                "inspection_id": inspection.inspection_id,
                "cv_analysis": cv_analysis,
                "ai_analysis": ai_analysis,
                "defects_found": len(inspection.defects_detected),
                "message": "Visual inspection completed successfully" if success else "Error processing visual inspection"
            }
            
        except Exception as e:
            logger.error(f"Error processing visual inspection: {e}")
            return {"success": False, "error": str(e)}
    
    async def start_diagnosis(self, machine_id: str, user_description: str = "") -> Dict[str, Any]:
        """Start comprehensive diagnosis"""
        try:
            # Start interactive session
            session = await self.diagnostic_engine.interactive_diagnosis_session(machine_id)
            
            # Perform initial comprehensive diagnosis
            diagnosis = await self.diagnostic_engine.comprehensive_diagnosis(
                machine_id=machine_id,
                user_description=user_description
            )
            
            # Update session with diagnosis
            session.current_diagnosis = diagnosis
            self.knowledge_base.save_diagnosis_session(session)
            
            return {
                "success": True,
                "session_id": session.session_id,
                "diagnosis": {
                    "fault_id": diagnosis.fault_id,
                    "severity": diagnosis.severity,
                    "confidence": diagnosis.confidence,
                    "root_cause": diagnosis.root_cause,
                    "recommended_actions": diagnosis.recommended_actions,
                    "estimated_downtime": diagnosis.estimated_downtime
                },
                "message": "Diagnosis session started successfully"
            }
            
        except Exception as e:
            logger.error(f"Error starting diagnosis: {e}")
            return {"success": False, "error": str(e)}
    
    async def chat_with_diagnosis(self, session_id: str, message: str, uploaded_files: Optional[List[str]] = None) -> Dict[str, Any]:
        """Continue diagnosis conversation"""
        try:
            response = await self.diagnostic_engine.continue_diagnosis_conversation(
                session_id=session_id,
                user_message=message,
                uploaded_files=uploaded_files or []
            )
            
            return {
                "success": True,
                "response": response,
                "session_id": session_id
            }
            
        except Exception as e:
            logger.error(f"Error in diagnosis chat: {e}")
            return {"success": False, "error": str(e)}
    
    async def get_machine_overview(self, machine_id: str) -> Dict[str, Any]:
        """Get comprehensive machine overview"""
        try:
            # Get machine documents
            documents = self.knowledge_base.get_machine_documents(machine_id)
            
            # Get visual inspections
            inspections = self.knowledge_base.get_visual_inspections(machine_id)
            
            # Calculate summary statistics
            total_defects = sum(len(insp.defects_detected) for insp in inspections)
            recent_inspections = [insp for insp in inspections if insp.timestamp > datetime.now() - timedelta(days=30)]
            
            return {
                "success": True,
                "machine_id": machine_id,
                "documentation": {
                    "total_documents": len(documents),
                    "document_types": list(set(doc.file_type for doc in documents)),
                    "latest_upload": max([doc.upload_date for doc in documents]).isoformat() if documents else None
                },
                "visual_inspections": {
                    "total_inspections": len(inspections),
                    "recent_inspections": len(recent_inspections),
                    "total_defects_detected": total_defects,
                    "latest_inspection": max([insp.timestamp for insp in inspections]).isoformat() if inspections else None
                },
                "status": "Active" if documents or inspections else "No data available"
            }
            
        except Exception as e:
            logger.error(f"Error getting machine overview: {e}")
            return {"success": False, "error": str(e)}
    
    async def close(self):
        """Cleanup system resources"""
        await self.llm_provider.close()

# Example usage and testing
async def main():
    """Example usage of the Enhanced Smart Maintenance System"""
    system = EnhancedSmartMaintenanceSystem()
    
    try:
        # Test Gemini connection
        print("Testing Gemini API connection...")
        test_response = await system.llm_provider.call_gemini("What are the main causes of bearing failure in industrial motors?")
        print(f"Gemini Test Response: {test_response[:200]}...")
        
        # Add some sample equipment
        sample_equipment = {
            'id': 'PUMP_001',
            'name': 'Main Water Pump',
            'type': 'Centrifugal Pump',
            'model': 'CP-500',
            'manufacturer': 'Industrial Pumps Inc.',
            'installation_date': '2020-01-15',
            'last_maintenance': '2024-01-15'
        }
        
        system.knowledge_base.add_equipment(sample_equipment)
        print("Sample equipment added")
        
        # Example: Start diagnosis
        diagnosis_result = await system.start_diagnosis(
            machine_id='PUMP_001',
            user_description="Pump is making unusual noise and vibrating more than normal"
        )
        print(f"Gemini Diagnosis started: {diagnosis_result}")
        
        # Example: Chat with diagnosis system
        if diagnosis_result.get('success'):
            chat_result = await system.chat_with_diagnosis(
                session_id=diagnosis_result['session_id'],
                message="The noise seems to be coming from the bearing area. What should I check first?"
            )
            print(f"Gemini Diagnosis chat response: {chat_result}")
        
        # Example: Get machine overview
        overview = await system.get_machine_overview('PUMP_001')
        print(f"Machine overview: {overview}")
        
    except Exception as e:
        print(f"Error in main: {e}")
    
    finally:
        await system.close()

if __name__ == "__main__":
    asyncio.run(main())